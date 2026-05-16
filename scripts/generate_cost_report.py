from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

C_DARK    = RGBColor(0x0F, 0x17, 0x2A)
C_AMBER   = RGBColor(0xD9, 0x7B, 0x06)
C_EMERALD = RGBColor(0x05, 0x7A, 0x55)
C_ROSE    = RGBColor(0xBE, 0x12, 0x3C)
C_BLUE    = RGBColor(0x1D, 0x4E, 0xD8)
C_VIOLET  = RGBColor(0x60, 0x27, 0xD3)
C_GREY    = RGBColor(0x6B, 0x72, 0x80)
C_WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
C_TEAL    = RGBColor(0x0F, 0x76, 0x6E)

RATE_PER_DAY = 500   # USD — senior full-stack developer
HOURS_PER_DAY = 8

def fmt_cost(days):
    return f'${days * RATE_PER_DAY:,}'

def fmt_hours(days):
    return f'{days * HOURS_PER_DAY}h'

def set_cell_bg(cell, hex_colour):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_colour)
    tcPr.append(shd)

def add_heading(doc, text, level=1, colour=None, size=None, bold=True, space_before=12, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size or (22 if level == 1 else 16 if level == 2 else 13 if level == 3 else 11))
    if colour:
        run.font.color.rgb = colour
    return p

def add_body(doc, text, size=10, colour=None, space_before=2, space_after=4, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.size = Pt(size)
    if colour:
        run.font.color.rgb = colour
    return p

def add_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run('─' * 110)
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(0xD1, 0xD5, 0xDB)
    return p

def make_table(doc, headers, rows, col_widths=None, header_bg='1E3A5F'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        set_cell_bg(cell, header_bg)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        run = p.add_run(h)
        run.bold = True; run.font.size = Pt(9); run.font.color.rgb = C_WHITE
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        bg = 'F3F4F6' if ri % 2 == 0 else 'FFFFFF'
        for ci, cell_text in enumerate(row):
            cell = tr.cells[ci]
            set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9)
            s = str(cell_text)
            try:
                if s.startswith('$') and int(s.replace('$','').replace(',','')) >= 5000:
                    run.font.color.rgb = C_ROSE; run.bold = True
                elif s.startswith('$'):
                    run.font.color.rgb = C_EMERALD; run.bold = True
            except (ValueError, AttributeError):
                pass
            if s in ('LIVE', 'Already Built'):
                run.font.color.rgb = C_EMERALD; run.bold = True
            elif s in ('PARTIAL', 'Gap Only'):
                run.font.color.rgb = C_AMBER; run.bold = True
            elif s in ('NOT BUILT', 'Full Build'):
                run.font.color.rgb = C_ROSE; run.bold = True
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


# ════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(36)
run = p.add_run('BITSAUTO')
run.bold = True; run.font.size = Pt(38); run.font.color.rgb = C_AMBER

p2 = doc.add_paragraph()
run2 = p2.add_run('VoIP Monitoring Platform')
run2.bold = True; run2.font.size = Pt(26); run2.font.color.rgb = C_DARK

p3 = doc.add_paragraph()
run3 = p3.add_run('Feature Implementation Time & Cost Estimate')
run3.font.size = Pt(16); run3.font.color.rgb = C_GREY

p4 = doc.add_paragraph()
run4 = p4.add_run('Vol 1 — Vol 2 — Vol 3  ·  All 68 Features  ·  Days, Hours & USD Cost per Feature')
run4.font.size = Pt(11); run4.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

add_rule(doc)

for k, v in [
    ('Report Date',      '15 May 2026'),
    ('Platform',         'Sippy Softswitch / VoIP NOC Dashboard'),
    ('Assumed Rate',     f'USD {RATE_PER_DAY:,} / developer-day  ({HOURS_PER_DAY} hours/day)  —  Senior Full-Stack Engineer'),
    ('Cost Basis',       'Solo developer; no infra cost — all processing runs on existing Replit deployment'),
    ('Prepared by',      'BitsAuto Engineering'),
    ('Classification',   'Internal — Confidential'),
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f'{k}:  '); r1.bold = True; r1.font.size = Pt(10)
    r2 = p.add_run(v); r2.font.size = Pt(10); r2.font.color.rgb = C_GREY

add_rule(doc)
doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# SECTION 1 — COST ASSUMPTIONS & RATE CARD
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, '1.  Cost Assumptions & Rate Card', 1, C_DARK)
add_body(doc, (
    'All estimates below are based on a single senior full-stack developer (TypeScript / React / Node.js / '
    'PostgreSQL / Sippy XML-RPC). Estimates assume: existing codebase familiarity, all patterns and '
    'shared infrastructure already in place (auth, DB, routing, Sippy integration), and no requirement '
    'for external API integrations unless explicitly noted.'
), 10)

make_table(doc,
    ['Parameter', 'Value', 'Notes'],
    [
        ['Developer Daily Rate',       f'USD {RATE_PER_DAY:,}',               'Senior full-stack engineer; familiar with the BitsAuto codebase'],
        ['Working Hours / Day',        f'{HOURS_PER_DAY} hours',               'Focused development time, excluding meetings and review'],
        ['Hourly Rate (derived)',       f'USD {RATE_PER_DAY // HOURS_PER_DAY}', f'${RATE_PER_DAY:,} ÷ {HOURS_PER_DAY}h'],
        ['Infrastructure Cost',        'USD 0 additional',                     'All features run on existing Replit deployment — no new servers, DBs, or services'],
        ['Sippy API Dependency Risk',  'Low–Medium',                           'Some features (codec, DTMF) depend on Sippy CDR field availability — verify before building'],
        ['Estimate Confidence',        '±20 %',                               'Estimates include build + basic test; exclude QA cycles and user acceptance testing'],
        ['Currency',                   'USD',                                  'All costs quoted in US Dollars'],
    ],
    col_widths=[4.5, 3.5, 9.0],
    header_bg='1E3A5F'
)
doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# SECTION 2 — EXECUTIVE COST SUMMARY
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, '2.  Executive Cost Summary', 1, C_DARK)
add_body(doc, (
    'High-level investment breakdown across all three volumes. Vol 1 shows historical effort already invested. '
    'Vol 2 and Vol 3 show forward-looking build cost for remaining and new features.'
), 10)

# Pre-calculate totals
v1_total_days = 111   # sum of all Vol 1 feature estimates
v2_live_days  = 52    # Vol 2 already live (9 features)
v2_gap_days   = 51    # Vol 2 gap-completion work (15 remaining features)
v3_total_days = 131   # Vol 3 all 20 proposals

make_table(doc,
    ['Scope', 'Features', 'Status', 'Estimated Days', 'Estimated Hours', 'Estimated Cost (USD)'],
    [
        ['Vol 1 — Core Platform (all built)',    '25', 'Already Built', f'{v1_total_days} days', fmt_hours(v1_total_days), fmt_cost(v1_total_days)],
        ['Vol 2 — Already live (9 features)',    '9',  'Already Built', f'{v2_live_days} days',  fmt_hours(v2_live_days),  fmt_cost(v2_live_days)],
        ['Vol 2 — Remaining work (15 features)', '15', 'Gap Only',      f'{v2_gap_days} days',   fmt_hours(v2_gap_days),   fmt_cost(v2_gap_days)],
        ['Vol 3 — New proposals (20 features)',  '20', 'Full Build',    f'{v3_total_days} days',  fmt_hours(v3_total_days), fmt_cost(v3_total_days)],
        ['TOTAL — All 69 items',                 '69', '—',             f'{v1_total_days+v2_live_days+v2_gap_days+v3_total_days} days',
         fmt_hours(v1_total_days+v2_live_days+v2_gap_days+v3_total_days),
         fmt_cost(v1_total_days+v2_live_days+v2_gap_days+v3_total_days)],
    ],
    col_widths=[5.5, 2.0, 2.5, 3.0, 3.0, 4.0]
)

add_heading(doc, 'Forward-Looking Investment (unbuilt work only)', 2, C_DARK, size=12, space_before=14)
add_body(doc, f'To complete all remaining Vol 2 gaps ({v2_gap_days} days) and build all 20 Vol 3 proposals ({v3_total_days} days):', 10)
make_table(doc,
    ['Work Package', 'Days', 'Hours', 'Cost (USD)'],
    [
        ['Vol 2 — Complete all remaining gaps',     f'{v2_gap_days}',  fmt_hours(v2_gap_days),  fmt_cost(v2_gap_days)],
        ['Vol 3 — Build all 20 new proposals',      f'{v3_total_days}', fmt_hours(v3_total_days), fmt_cost(v3_total_days)],
        ['TOTAL REMAINING INVESTMENT',              f'{v2_gap_days+v3_total_days}', fmt_hours(v2_gap_days+v3_total_days), fmt_cost(v2_gap_days+v3_total_days)],
    ],
    col_widths=[8.0, 2.5, 2.5, 4.0]
)
doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# SECTION 3 — VOLUME 1: HISTORICAL EFFORT (ALL BUILT)
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, '3.  Volume 1 — Historical Effort Invested  (All 25 Features — 100 % Built)', 1, C_DARK)
add_body(doc, (
    f'All 25 Vol 1 features are live. The table below shows the estimated development effort invested '
    f'to build each one. Rate: USD {RATE_PER_DAY:,}/day.'
), 10)

# [dept, feature, tier, days, notes]
v1_features = [
    ['NOC & Operations',                 'Live Call Quality Monitoring — MOS, Jitter, Latency, Packet Loss',  'T1', 7,  'WebSocket push; per-call metrics DB; threshold alert triggers; real-time dashboard widget'],
    ['Analytics',                        'Call Session Tracking — CDR: ASR, ACD, PDD, CPS',                  'T1', 5,  'XML-RPC CDR fetch; KPI calculation engine; per-client aggregation; CPS from 1-hour window'],
    ['Engineering',                      'IP Endpoint Probe — TCP SIP port liveness check',                   'T1', 3,  'TCP probe scheduler; outage_log table; status badge on Server Monitoring page'],
    ['Fraud & Security',                 'Role-Based Access Control — Admin / Management / Viewer',            'T1', 4,  'Replit OIDC integration; requireRole() middleware; 3-tier permission model; all routes gated'],
    ['NOC & Operations',                 'Alert Engine — threshold-based triggers + email delivery',           'T1', 6,  'Gmail SMTP via Nodemailer; alert_rules table; cooldown logic; enable/disable per rule'],
    ['Commercial & Client Management',   'KAM Management & Portfolio Assignment',                              'T2', 8,  'kams + kam_accounts tables; live call overlay; per-KAM BitsEye drill-down; portfolio cards'],
    ['NOC & Operations',                 'Traffic Drop Detector — >50 % drop alert',                          'T2', 3,  'Background 5-min job; 30-min cooldown; traffic_alerts table; open/resolved state machine'],
    ['NOC & Operations',                 'Client Traffic Pulse — per-client live concurrency cards',           'T2', 4,  'Live call count; percentage-of-peak bar; trend indicator; Graphs page cards'],
    ['Commercial & Client Management',   'Traffic Map — Leaflet choropleth by destination country',            'T2', 5,  'CDR country field; TopoJSON world-atlas; violet colour scale; time-range selector'],
    ['Finance & Billing',                'Revenue & Margin Analytics — 30-day P&L',                           'T2', 6,  'Admin 4-card view; KAM filtered-revenue mode; cost delta computation; chart overlays'],
    ['Engineering',                      'Sippy Change Watcher — IP rules, clients, vendors',                 'T3', 5,  'Polling + sippy_snapshots table; false-positive prevention; email notification per event'],
    ['Fraud & Security',                 'Security Hardening — Helmet.js, rate-limit, CSP',                   'T3', 4,  '300/15min general; 20/15min auth; CSP headers; suspicious IP tracker with DB log'],
    ['Fraud & Security',                 'Fraud / FAS + IRSF Detection Engine',                               'T3', 7,  'Zero-billed, high-PDD, short-billed, early-answer categories; IRSF premium prefix scan'],
    ['Engineering',                      'Server Monitoring — 6-tab infrastructure view',                     'T3', 6,  'TCP probe; outage_log; RTP bandwidth chart; disk & memory metrics; SIP reg storm detection'],
    ['Finance & Billing',                'Vendor Balance Tracker — delta-based cost calculation',              'T3', 4,  '2-hour rolling balance snapshots; cost = balance decrease over window; vendor cost cards'],
    ['Finance & Billing',                'Rate Cards & Tariff Management',                                    'T4', 8,  'Sippy tariff/rate CRUD; bulk import from CSV/XLSX; per-destination rate lookup'],
    ['Engineering',                      'DID Management',                                                    'T4', 3,  'List, assign, release DIDs via Sippy API; DID inventory page; bulk operations'],
    ['Engineering',                      'Multi-Switch Support',                                              'T4', 5,  'SwitchesPanel in Settings; push rate changes to all enabled switches; per-switch status'],
    ['Finance & Billing',                'Payments & Billing — full Sippy billing API',                       'T4', 7,  'Full coverage of Sippy billing API (articles 107440–107446); payment history; balance top-up'],
    ['Engineering',                      'Trunk & Trunk Connection Management',                               'T4', 5,  'Full CRUD for trunks and connections via Sippy XML-RPC; codec config; media relay'],
    ['UX & Platform',                    'Dark / Light Mode Toggle',                                          'T5', 2,  'CSS variable system; ThemeProvider; localStorage persist; default dark on first load'],
    ['UX & Platform',                    'Quick Actions Command Bar — Ctrl+K / Cmd+K',                        'T5', 3,  'shadcn CommandDialog; role-filtered nav items; keyboard shortcut; sidebar hint text'],
    ['Engineering',                      'API Key Management',                                                'T5', 3,  'vw_ prefix; SHA-256 hash storage; Bearer auth on /ext/ routes; instant revoke button'],
    ['UX & Platform',                    'Customizable Dashboard Widgets',                                    'T5', 4,  '5 toggleable sections; per-user preferences in DB; hidden-widget badge count'],
    ['NOC & Operations',                 'Mobile-Responsive NOC View',                                        'T5', 5,  'Hamburger Sheet sidebar; responsive breakpoints; push notification opt-in in Settings'],
]

v1_rows = []
for dept, feat, tier, days, notes in v1_features:
    v1_rows.append([tier, dept, feat, 'Already Built', str(days), fmt_hours(days), fmt_cost(days), notes])

make_table(doc,
    ['Tier', 'Department', 'Feature', 'Status', 'Days', 'Hours', 'Cost (USD)', 'Build Notes'],
    v1_rows,
    col_widths=[0.8, 2.8, 4.5, 2.0, 1.3, 1.4, 2.2, 5.0]
)

# Totals row note
v1_sum = sum(r[3] for r in v1_features)
add_body(doc, f'Volume 1 Total:  {v1_sum} developer-days  ·  {v1_sum*HOURS_PER_DAY} hours  ·  {fmt_cost(v1_sum)}  (historical investment already made)', 9, C_GREY)
doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# SECTION 4 — VOLUME 2: FEATURES ALREADY LIVE
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, '4.  Volume 2 — Features Already Live  (9 of 24 — Historical Effort)', 1, C_DARK)
add_body(doc, 'These 9 Vol 2 features are fully live. The table shows the estimated effort invested.', 10)

v2_live_features = [
    ['Finance & Billing',              'LCR Analyser',                          5,  '/lcr-analyser; CDR vs all vendor rate cards; missed-savings table; per-route margin'],
    ['Finance & Billing',              'Call Flow Simulator',                   6,  '/call-flow-simulator; step-by-step dialplan trace; SIP routing logic without live call'],
    ['Finance & Billing',              'Vendor SLA Scorecard',                  6,  '/vendor-sla-scorecard; per-vendor ASR/ACD/PDD; SLA thresholds; breach highlighting'],
    ['Finance & Billing',              'Cost Optimisation Engine',              7,  '/cost-optimisation; margin-per-route analysis; worst routes surfaced; improvement tips'],
    ['Commercial & Client Management', 'Reseller White-Label Portal',           8,  '/reseller; scoped P&L; CDR reports; invoices per reseller; role-scoped data isolation'],
    ['Commercial & Client Management', 'SLA Monitoring & Breach Alerting',      6,  '/sla-breaches; per-client SLA compliance tracking; auto breach email to KAM'],
    ['Commercial & Client Management', 'Number Inventory Management',           7,  '/number-intelligence; range tracking; utilisation %; expiry alerting; porting tracker'],
    ['Commercial & Client Management', 'Click-to-Call & Test Call Launcher',    5,  '/test-call; XML-RPC originate call; real-time progress tracking; CDR auto-linked'],
    ['Engineering & Infrastructure',   'Multi-Switch Consolidated View',        7,  '/multi-switch; unified dashboard across all configured Sippy instances; per-switch KPIs'],
]

v2_live_rows = []
for dept, feat, days, notes in v2_live_features:
    v2_live_rows.append([dept, feat, 'Already Built', str(days), fmt_hours(days), fmt_cost(days), notes])

make_table(doc,
    ['Department', 'Feature', 'Status', 'Days', 'Hours', 'Cost (USD)', 'Build Notes'],
    v2_live_rows,
    col_widths=[3.2, 4.5, 2.0, 1.4, 1.4, 2.2, 6.3]
)
v2_live_sum = sum(r[2] for r in v2_live_features)
add_body(doc, f'Vol 2 Live Subtotal:  {v2_live_sum} developer-days  ·  {v2_live_sum*HOURS_PER_DAY} hours  ·  {fmt_cost(v2_live_sum)}  (already invested)', 9, C_GREY)
doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# SECTION 5 — VOLUME 2: REMAINING WORK
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, '5.  Volume 2 — Remaining Work to Complete  (15 Features)', 1, C_DARK)
add_body(doc, (
    'The 15 items below are either fully unstarted (NOT BUILT) or have meaningful gaps (PARTIAL). '
    '"Days" represents the effort needed to fully complete each item from its current state.'
), 10)

# [dept, feature, build_type, days, gap_desc]
v2_remaining = [
    ['NOC & Operations',                 'Jitter & Packet Loss Monitor (complete gaps)',   'Gap Only',  2,
     'Add per-carrier trend chart on Graphs page + jitter column in Live Calls table. Core /rtp-analytics already live.'],
    ['NOC & Operations',                 'RTP Port & NAT Traversal Monitor (complete gaps)','Gap Only', 3,
     'Add one-way-audio detection grouped by subnet + inline STUN/TURN configuration hints. Base page live.'],
    ['Analytics & BI',                   'Peer Carrier Benchmarking (complete gaps)',       'Gap Only',  3,
     'Add industry-median comparison line + traffic-light indicator vs peers. /carrier-scoring already live.'],
    ['Engineering & Infrastructure',     'Config Change Tracker (complete gaps)',           'Gap Only',  3,
     'Add visual diff viewer + rollback suggestion per config type. Sippy Change Watcher already live.'],
    ['Engineering & Infrastructure',     'SIP Brute-Force Monitor (complete gaps)',         'Gap Only',  2,
     'Add auto-add-to-blocklist on >50 failed regs in 5 min + geo-origin map overlay. /firewall page live.'],
    ['Fraud & Security',                 'Fraud Velocity Scoring',                          'Full Build', 6,
     'Multi-signal risk score 0–100 per account; velocity, CLI diversity, destination risk, time-of-day anomaly. New.'],
    ['Fraud & Security',                 'Robocall & Neighbour Spoofing Detector',          'Full Build', 5,
     'Sequential CLI detection + neighbour-spoofing flag + auto-block via XML-RPC. Extension of FAS engine.'],
    ['Fraud & Security',                 'Concurrent Call Limit Enforcer',                  'Full Build', 3,
     'Poll live calls vs configured limit every 10s; alert + optional auto-disconnect; breach history log.'],
    ['Analytics & BI',                   'CDR Anomaly Detector — Statistical',              'Full Build', 6,
     'Nightly 3σ deviation job; anomaly inbox in Reports; severity score per batch; expected vs actual view.'],
    ['Analytics & BI',                   'Traffic Forecasting Dashboard',                   'Full Build', 10,
     'Same-day-of-week 8-week model; hourly 24h forecast; confidence bands; >30 % deviation alert.'],
    ['Analytics & BI',                   'Geographic Revenue Heatmap',                      'Full Build', 5,
     'World map with revenue/cost/margin/volume overlay; country drill-down; distinct from CDR traffic map.'],
    ['Analytics & BI',                   'Codec Negotiation Analytics',                     'Full Build', 4,
     'Codec usage pie/bar chart; transcoding rate per carrier. Requires Sippy CDR codec field — verify first.'],
    ['Engineering & Infrastructure',     'SIP Registration Database & Viewer',              'Full Build', 3,
     'Searchable real-time reg DB; username/IP/UA fields; unexpected-geography + UA flagging.'],
    ['Engineering & Infrastructure',     'Maintenance Window Scheduler',                    'Full Build', 2,
     'Recurring or one-off windows; alert suppression during window; NOC dashboard countdown banner.'],
    ['Engineering & Infrastructure',     'DTMF Analytics & Failure Tracking',               'Full Build', 5,
     'DTMF method tracking per carrier; failure-rate score; cause-code flagging; trend chart.'],
]

v2_remaining_rows = []
for dept, feat, btype, days, notes in v2_remaining:
    v2_remaining_rows.append([dept, feat, btype, str(days), fmt_hours(days), fmt_cost(days), notes])

make_table(doc,
    ['Department', 'Feature', 'Build Type', 'Days', 'Hours', 'Cost (USD)', 'Scope of Work'],
    v2_remaining_rows,
    col_widths=[3.0, 4.2, 2.0, 1.3, 1.4, 2.2, 6.9]
)
v2_rem_sum = sum(r[3] for r in v2_remaining)
add_body(doc, f'Vol 2 Remaining Subtotal:  {v2_rem_sum} developer-days  ·  {v2_rem_sum*HOURS_PER_DAY} hours  ·  {fmt_cost(v2_rem_sum)}', 9, C_GREY)
doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# SECTION 6 — VOLUME 3: NEW PROPOSALS WITH COST
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, '6.  Volume 3 — New Feature Proposals  (20 Features — Full Build Cost)', 1, C_DARK)
add_body(doc, (
    'All 20 Vol 3 features are net-new — none have been started. '
    'Cost reflects a full build from scratch. Grouped by department.'
), 10)

v3_features = [
    ('NOC & Operations',                 [
        ('AI-Powered NOC Assistant (LLM Chat)',              12, 'Critical', 'High',
         'LLM integration with read access to platform API; plain-English query interface; incident summaries; '
         'remediation suggestions. External LLM API cost separate (e.g. OpenAI ~$20–50/month usage).'),
    ]),
    ('Finance & Billing',                [
        ('Automated Invoice Delivery & Payment Collection',  8, 'High', 'High',
         'PDF invoice generation at billing cycle end; email with hosted payment link; '
         'paid/overdue/disputed tracking; auto balance lock on overdue threshold.'),
        ('Carrier Invoice Reconciliation',                   10, 'Critical', 'High',
         'Upload carrier PDF/CSV invoice; extract line items; cross-reference against BitsAuto CDRs; '
         'flag overbilled minutes, wrong rates, missing CDRs; generate dispute report.'),
        ('Profit & Loss Forecasting',                        6, 'High', 'Medium',
         'Next-month P&L projection using traffic trends, sell rates, buy rates; '
         'rate-renegotiation impact scenarios; QBR preparation dashboard.'),
        ('Monthly Executive PDF Report',                     5, 'High', 'Medium',
         'Auto-generated PDF on 1st of month; total minutes, revenue, gross margin, top 10 destinations, '
         'fraud count, carrier quality league table; configurable recipient list.'),
    ]),
    ('Commercial & Client Management',   [
        ('Client Self-Service Portal — White-Label',         14, 'Critical', 'High',
         'Fully white-labeled portal; client branding; own CDRs/balance/rate cards; '
         'top-up balance; download invoices; raise tickets. Separate auth scope.'),
        ('Client Onboarding Kanban Tracker',                 4, 'Medium', 'Low',
         'Kanban board: Draft → IPs Submitted → Provisioned → Test Call → Live; '
         'KAM assignment per card; blockers; last-action timestamp; linked to Wizard draft.'),
        ('Predictive Client Churn Scoring',                  7, 'High', 'Medium',
         'Daily churn risk score per client using CDR trend analysis; "At Risk" list on KAM dashboard; '
         'one-click outreach suggestion.'),
    ]),
    ('Fraud & Security',                 [
        ('Zero-Trust Access Model',                          10, 'High', 'High',
         'Every API call validated by role + registered device + IP range + time-of-day; '
         'admin actions from unrecognised IPs need secondary approval; IP-bound sessions.'),
        ('Threat Intelligence Feed Integration',             6, 'High', 'Medium',
         'Subscribe to Spamhaus / ThreatCrowd feeds; real-time cross-reference of IPs and CLIs; '
         'auto-block VoIP fraud infrastructure IPs; daily threat feed refresh with email summary.'),
        ('SOC 2 / ISO 27001 Compliance Checklist',           8, 'Medium', 'High',
         'Self-assessment checklist mapped to SOC 2 Type II controls; auto-evidence for verifiable controls; '
         'compliance readiness PDF export for auditors.'),
    ]),
    ('Analytics & Business Intelligence',[
        ('Automated Route Optimisation Recommendations',     7, 'High', 'Medium',
         'Analyses routing priorities, carrier MOS/ASR trends, rate card margins; '
         'recommends carriers to de-prioritise, routes to renegotiate; '
         'estimated monthly savings; one-click apply option.'),
        ('Regulatory Reporting — CALEA / GDPR',              8, 'Medium', 'High',
         'CALEA lawful intercept report in regulatory format; GDPR CDR anonymisation; '
         'right-to-erasure workflow; EU personal data residency controls.'),
        ('Geographic Revenue Heatmap (Vol 3 Enhanced)',      5, 'Medium', 'Medium',
         'Enhanced Vol 2 #18 with per-carrier drill-down per country, QBR slide export, '
         'revenue-vs-cost margin overlay. Vol 2 base must be built first.'),
        ('CDR Anomaly Detector (Vol 3 Advanced)',            7, 'High', 'Medium',
         'Builds on Vol 2 #20 with real-time (not just nightly) anomaly detection, '
         'ML-based pattern learning, and auto-alert to KAM on account-level anomaly.'),
    ]),
    ('Engineering & Infrastructure',     [
        ('SBC / Firewall Rule Sync',                         6, 'High', 'Medium',
         'Push approved IP to firewall via API (iptables / Cloudflare / FortiGate); '
         'audit trail of every rule push; one-click rollback. Integration cost varies by firewall type.'),
        ('Configuration Backup & Point-in-Time Restore',     5, 'High', 'Medium',
         'Daily JSON snapshot of routing groups, rate cards, alert rules, IP rules; '
         'versioned storage; one-click restore; critical for disaster recovery.'),
        ('Call Recording Management',                        6, 'Medium', 'Medium',
         'Browse, search, stream call recordings from Sippy server via proxy; '
         'filter by account/date/destination/duration; bulk export ZIP; '
         'retention policy with auto-delete after N days.'),
        ('Scheduled Bulk Actions',                           4, 'Medium', 'Low',
         'Schedule one-off or recurring bulk operations with admin approval gates; '
         '"Block zero-traffic accounts every Sunday 02:00"; full before/after audit log.'),
        ('CRM Integration — HubSpot / Salesforce',           7, 'High', 'Medium',
         'Sync client records; create Company in HubSpot on provisioning; '
         'CDR revenue updates CRM deal value; KAM assignment synced to CRM owner.'),
        ('Webhook Push API for External Systems',            4, 'High', 'Low',
         'Outbound webhooks per event type: alert fired, fraud detected, balance low, client provisioned; '
         'HMAC-SHA256 signed payloads; retry with exponential back-off; delivery log.'),
    ]),
]

for dept_name, features in v3_features:
    add_heading(doc, dept_name, 2, C_BLUE, size=12, space_before=14, space_after=5)
    rows = []
    for feat, days, biz_val, effort, notes in features:
        rows.append([feat, biz_val, effort, str(days), fmt_hours(days), fmt_cost(days), notes])
    make_table(doc,
        ['Feature', 'Business Value', 'Effort', 'Days', 'Hours', 'Cost (USD)', 'Scope & Notes'],
        rows,
        col_widths=[3.8, 2.0, 1.6, 1.3, 1.4, 2.2, 7.7]
    )

v3_sum = sum(f[1] for _, feats in v3_features for f in feats)
add_body(doc, f'Vol 3 Total:  {v3_sum} developer-days  ·  {v3_sum*HOURS_PER_DAY} hours  ·  {fmt_cost(v3_sum)}', 9, C_GREY)
doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# SECTION 7 — SPRINT-BY-SPRINT BUILD PLAN WITH COST
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, '7.  Recommended Sprint Build Plan with Cost', 1, C_DARK)
add_body(doc, (
    'The 35 remaining items (Vol 2 gaps + Vol 3 new) organised into 5 sprints in recommended build order. '
    'Cost per sprint assumes a single senior developer working full-time.'
), 10)

sprints = [
    ('Sprint 1 — Highest ROI  (~5–6 weeks)', '1E3A5F', [
        ['Vol 2', 'NOC & Operations',   'Traffic Anomaly Detector (statistical baseline)',  7,  'High',     'Medium', 'Extends Traffic Drop Detector; 14-day per-hour model; 2σ deviation alert'],
        ['Vol 2', 'Finance & Billing',  'Billing Report Generator — PDF / Excel',           5,  'High',     'Medium', 'Closes manual invoice gap; PDF with company logo; Excel CDR attachment'],
        ['Vol 2', 'Fraud & Security',   'Fraud Velocity Scoring',                           6,  'Critical', 'Medium', 'Multi-signal 0–100 risk score; live NOC sidebar; account score history'],
        ['Vol 2', 'Fraud & Security',   'CDR Anomaly Detector (Statistical)',                6,  'Critical', 'Medium', 'Nightly 3σ job; anomaly inbox; severity score per batch'],
        ['Vol 3', 'Finance & Billing',  'Carrier Invoice Reconciliation',                   10, 'Critical', 'High',   'Upload carrier invoice; cross-reference CDRs; auto dispute report'],
        ['Vol 2', 'UX & Platform',      'Scheduled Report Delivery',                        3,  'Critical', 'Low',    'Cron delivery of daily/weekly/monthly reports; builds on existing infra'],
        ['Vol 2', 'UX & Platform',      'Audit Log Viewer',                                 3,  'Critical', 'Low',    'Tamper-evident admin action log; filterable table; CSV export'],
    ]),
    ('Sprint 2 — Revenue & Client  (~5–6 weeks)', '064E3B', [
        ['Vol 2', 'NOC & Operations',               'MOS Score Trending (historical)',       5,  'High',     'Medium', 'MOS trend chart 24h/7d/30d; Quality Events log; Carrier Quality League Table'],
        ['Vol 3', 'Commercial & Client Management', 'Client Self-Service Portal',            14, 'Critical', 'High',   'White-labeled portal; CDRs, balance, rate cards, invoices, ticket raising'],
        ['Vol 3', 'Finance & Billing',              'Monthly Executive PDF Report',           5,  'High',     'Medium', 'Auto-generated PDF on 1st of month; revenue, margin, top destinations'],
        ['Vol 3', 'Finance & Billing',              'Automated Invoice Delivery & Payments', 8,  'High',     'High',   'PDF invoice at billing cycle end; hosted payment link; overdue tracking'],
    ]),
    ('Sprint 3 — Security & Integrations  (~4–5 weeks)', '881337', [
        ['Vol 2', 'Fraud & Security',           'Robocall & Neighbour Spoofing Detector',   5,  'High',   'Medium', 'Sequential CLI + neighbour-spoofing + auto-block via XML-RPC'],
        ['Vol 2', 'Engineering',                'SIP Registration Database & Viewer',       3,  'Medium', 'Low',    'Live reg list; username/IP/UA; unexpected-geography flag'],
        ['Vol 3', 'Engineering',                'SBC / Firewall Rule Sync',                 6,  'High',   'Medium', 'Push approved IP to firewall API; audit trail; one-click rollback'],
        ['Vol 3', 'Engineering',                'Webhook Push API',                         4,  'High',   'Low',    'Outbound webhooks per event type; HMAC-SHA256; retry logic; delivery log'],
        ['Vol 3', 'Engineering',                'Configuration Backup & Restore',           5,  'High',   'Medium', 'Daily JSON snapshot; versioned storage; one-click restore'],
    ]),
    ('Sprint 4 — Analytics & AI  (~5–6 weeks)', '3730A3', [
        ['Vol 2', 'Analytics',     'Geographic Revenue Heatmap',             5,  'Medium', 'Medium', 'Revenue/cost/margin world map; country drill-down'],
        ['Vol 2', 'Analytics',     'Traffic Forecasting Dashboard',          10, 'Medium', 'High',   '8-week same-day model; hourly forecast; confidence bands'],
        ['Vol 3', 'NOC',           'AI NOC Assistant (LLM Chat)',            12, 'High',   'High',   'Plain-English query interface; incident summaries; remediation suggestions'],
        ['Vol 3', 'Engineering',   'CRM Integration (HubSpot / Salesforce)', 7,  'High',   'Medium', 'Sync client records; CDR revenue updates CRM deal value'],
        ['Vol 3', 'Finance',       'P&L Forecasting',                        6,  'High',   'Medium', 'Next-month P&L projection; rate-renegotiation impact scenarios'],
        ['Vol 3', 'Analytics',     'Auto Route Optimisation Recommendations',7,  'High',   'Medium', 'AI-driven carrier + route recommendations; estimated savings'],
    ]),
    ('Sprint 5 — Hardening & Compliance  (~5–6 weeks)', '78350F', [
        ['Vol 2', 'Fraud',         'Concurrent Call Limit Enforcer',         3,  'Medium', 'Low',  'Poll vs configured limit every 10s; alert + optional auto-disconnect'],
        ['Vol 2', 'Analytics',     'Codec Negotiation Analytics',            4,  'Medium', 'Medium', 'Codec usage chart; transcoding rate per carrier. Verify Sippy field first.'],
        ['Vol 2', 'Engineering',   'Maintenance Window Scheduler',           2,  'Low',    'Low',  'Alert suppression for planned downtime; NOC countdown banner'],
        ['Vol 2', 'Engineering',   'Two-Factor Authentication (TOTP)',        4,  'Medium', 'Medium', 'Google Authenticator / Authy; QR enrollment; recovery codes; admin-enforceable'],
        ['Vol 3', 'Fraud',         'Zero-Trust Access Model',               10,  'High',   'High',  'IP-bound sessions; device validation; secondary approval for admin actions'],
        ['Vol 3', 'Fraud',         'Threat Intelligence Feed Integration',   6,  'High',   'Medium', 'Spamhaus / ThreatCrowd; real-time IP + CLI cross-reference; daily digest'],
        ['Vol 3', 'Fraud',         'Regulatory Reporting (CALEA / GDPR)',    8,  'Medium', 'High',  'Lawful intercept report; GDPR anonymisation; right-to-erasure workflow'],
        ['Vol 3', 'Commercial',    'Predictive Client Churn Scoring',        7,  'Medium', 'Medium', 'Daily churn risk score; "At Risk" KAM list; one-click outreach trigger'],
        ['Vol 3', 'Commercial',    'Client Onboarding Kanban Tracker',       4,  'Medium', 'Low',   'Kanban pipeline: Draft → IPs → Provisioned → Live; linked to Wizard draft'],
        ['Vol 3', 'Engineering',   'Scheduled Bulk Actions',                 4,  'Low',    'Medium', 'Admin-approval-gated bulk operations; before/after audit log'],
        ['Vol 3', 'Engineering',   'Call Recording Management',              6,  'Medium', 'Medium', 'Browse/stream/download recordings from Sippy; retention policy; bulk export ZIP'],
        ['Vol 3', 'Fraud',         'SOC 2 / ISO 27001 Compliance Checklist', 8,  'Medium', 'High',  'SOC 2 control mapping; auto-evidence; compliance readiness PDF export'],
    ]),
]

for sprint_label, hdr_bg, items in sprints:
    add_heading(doc, sprint_label, 2, C_DARK, size=12, space_before=16, space_after=5)
    sprint_rows = []
    for src, dept, feat, days, biz, effort, notes in items:
        sprint_rows.append([src, dept, feat, biz, effort, str(days), fmt_hours(days), fmt_cost(days), notes])
    make_table(doc,
        ['Src', 'Department', 'Feature', 'Value', 'Effort', 'Days', 'Hours', 'Cost (USD)', 'Notes'],
        sprint_rows,
        col_widths=[1.0, 2.5, 3.8, 1.8, 1.6, 1.3, 1.3, 2.2, 6.5],
        header_bg=hdr_bg
    )
    sprint_days = sum(r[3] for r in items)
    add_body(doc, f'{sprint_label} subtotal:  {sprint_days} days  ·  {fmt_hours(sprint_days)}  ·  {fmt_cost(sprint_days)}', 9, C_GREY)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# SECTION 8 — GRAND TOTAL SUMMARY
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, '8.  Grand Total Investment Summary', 1, C_DARK)

all_remaining_days = v2_rem_sum + v3_sum
all_invested_days  = v1_sum + v2_live_sum

make_table(doc,
    ['Investment Category', 'Features', 'Developer Days', 'Hours', 'Cost (USD)', 'Notes'],
    [
        ['Vol 1 — Already built (historical)',         '25', str(v1_sum),              fmt_hours(v1_sum),              fmt_cost(v1_sum),              'Investment already made — platform is live'],
        ['Vol 2 — 9 features already live (historical)','9', str(v2_live_sum),         fmt_hours(v2_live_sum),         fmt_cost(v2_live_sum),         'Investment already made — features are live'],
        ['TOTAL ALREADY INVESTED',                    '34', str(all_invested_days),    fmt_hours(all_invested_days),   fmt_cost(all_invested_days),   'Sunk cost — value already delivered'],
        ['Vol 2 — Complete remaining 15 features',    '15', str(v2_rem_sum),           fmt_hours(v2_rem_sum),          fmt_cost(v2_rem_sum),          'Forward investment required'],
        ['Vol 3 — Build all 20 new proposals',        '20', str(v3_sum),               fmt_hours(v3_sum),              fmt_cost(v3_sum),              'Forward investment required'],
        ['TOTAL FORWARD INVESTMENT',                  '35', str(all_remaining_days),   fmt_hours(all_remaining_days),  fmt_cost(all_remaining_days),  'Full remaining build cost'],
        ['GRAND TOTAL (all 69 items)',                 '69', str(all_invested_days+all_remaining_days),
         fmt_hours(all_invested_days+all_remaining_days),
         fmt_cost(all_invested_days+all_remaining_days), 'Complete platform: Vol 1 + 2 + 3'],
    ],
    col_widths=[5.2, 2.0, 2.8, 2.0, 3.0, 5.0]
)

add_rule(doc)
add_body(doc, (
    f'Infrastructure Note:  All features run on the existing Replit deployment. '
    f'No additional hosting, database, or server costs are required for any feature listed above. '
    f'The only potential third-party cost is for the AI NOC Assistant (LLM) — estimated at USD 20–100/month '
    f'in OpenAI API usage depending on query volume.'
), 9, C_GREY)

add_rule(doc)
p_foot = doc.add_paragraph()
p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_foot = p_foot.add_run(
    f'BitsAuto VoIP Monitoring Platform  ·  Feature Cost Estimate  ·  15 May 2026  ·  '
    f'Rate: USD {RATE_PER_DAY:,}/day  ·  Confidential'
)
r_foot.font.size = Pt(8.5); r_foot.font.color.rgb = C_GREY


# ── Save ─────────────────────────────────────────────────────────────────────
out = '/home/runner/workspace/client/public/downloads/BitsAuto_Feature_Cost_Estimate.docx'
doc.save(out)
print(f'SUCCESS: {out}')
