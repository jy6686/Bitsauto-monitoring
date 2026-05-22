from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(1.9)
    section.bottom_margin = Cm(1.9)
    section.left_margin   = Cm(2.2)
    section.right_margin  = Cm(2.2)

# ── Bitsauto Brand Colour Palette ─────────────────────────────────────────────
C_NAVY      = RGBColor(0x0A, 0x16, 0x28)   # deepest bg
C_NAVY_MID  = RGBColor(0x0D, 0x20, 0x40)   # section banner bg
C_NAVY_CARD = RGBColor(0x11, 0x22, 0x40)   # card / info row bg
C_CYAN      = RGBColor(0x00, 0xD4, 0xFF)   # primary accent
C_CYAN_SOFT = RGBColor(0x38, 0xBD, 0xF8)   # secondary cyan
C_TEAL      = RGBColor(0x14, 0xB8, 0xA6)   # label accent
C_GREEN     = RGBColor(0x10, 0xB9, 0x81)   # LIVE / success
C_AMBER     = RGBColor(0xF5, 0x9E, 0x0B)   # PARTIAL / warning
C_ROSE      = RGBColor(0xF4, 0x3F, 0x5E)   # NOT BUILT / danger
C_GOLD      = RGBColor(0xFF, 0xD7, 0x00)   # numbers / highlights
C_WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
C_LIGHT     = RGBColor(0xE2, 0xE8, 0xF0)   # primary body text
C_SLATE     = RGBColor(0xCB, 0xD5, 0xE1)   # secondary body text
C_MID       = RGBColor(0x94, 0xA3, 0xB8)   # muted text
C_DIM       = RGBColor(0x47, 0x55, 0x69)   # very muted

# Hex strings for cell backgrounds
BG_NAVY      = '0A1628'
BG_NAVY_MID  = '0D2040'
BG_NAVY_CARD = '112240'
BG_DIVIDER   = '1E3A5F'

# Department colours — all updated to Bitsauto brand palette
DEPT_COLORS = {
    'NOC & Operations':                  (C_CYAN,      BG_NAVY_MID),
    'Finance & Billing':                 (C_GREEN,     '063B2A'),
    'Commercial & Client Management':    (C_CYAN_SOFT, '0D2A4A'),
    'Fraud & Security':                  (C_ROSE,      '3B0A1A'),
    'Analytics & Business Intelligence': (C_AMBER,     '3B2200'),
    'Engineering & Infrastructure':      (C_TEAL,      '053030'),
    'UX & Platform':                     (C_MID,       '1A2535'),
}

def set_cell_bg(cell, hex_colour):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_colour)
    tcPr.append(shd)

def add_heading(doc, text, level=1, colour=None, size=None, bold=True, space_before=14, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size or (22 if level == 1 else 16 if level == 2 else 13 if level == 3 else 11))
    if colour:
        run.font.color.rgb = colour
    return p

def add_body(doc, text, size=10, colour=None, space_before=2, space_after=5, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = colour or C_SLATE
    return p

def add_bullet(doc, text, size=10, colour=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = colour or C_SLATE
    return p

def add_rule(doc, colour=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run('─' * 110)
    run.font.size = Pt(7)
    run.font.color.rgb = colour or C_DIM
    return p

def add_section_banner(doc, title, subtitle=None):
    """Full-width dark navy banner with cyan title — Bitsauto brand style."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(0)
    r1 = p.add_run(f'  {title}')
    r1.bold = True
    r1.font.size = Pt(13)
    r1.font.color.rgb = C_CYAN
    # Shade the paragraph background
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), BG_NAVY_MID)
    pPr.append(shd)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after  = Pt(8)
        r2 = p2.add_run(f'  {subtitle}')
        r2.font.size = Pt(9)
        r2.italic = True
        r2.font.color.rgb = C_MID
        pPr2 = p2._p.get_or_add_pPr()
        shd2 = OxmlElement('w:shd')
        shd2.set(qn('w:val'), 'clear')
        shd2.set(qn('w:color'), 'auto')
        shd2.set(qn('w:fill'), BG_NAVY_MID)
        pPr2.append(shd2)
    else:
        doc.add_paragraph().paragraph_format.space_after = Pt(8)

def add_dept_banner(doc, dept_name):
    colour, bg = DEPT_COLORS[dept_name]
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(5)
    run = p.add_run(f'  ▌  {dept_name.upper()}')
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = colour
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), bg)
    pPr.append(shd)

def make_table(doc, headers, rows, col_widths=None, header_bg=BG_NAVY_MID):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        set_cell_bg(cell, header_bg)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = C_CYAN
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        bg = BG_NAVY_CARD if ri % 2 == 0 else BG_NAVY_MID
        for ci, cell_text in enumerate(row):
            cell = tr.cells[ci]
            set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9)
            s = str(cell_text)
            if s in ('LIVE', '✅ LIVE'):
                run.font.color.rgb = C_GREEN; run.bold = True
            elif s in ('PARTIAL', '⚠️ PARTIAL'):
                run.font.color.rgb = C_AMBER; run.bold = True
            elif 'NOT BUILT' in s or 'REMAINING' in s:
                run.font.color.rgb = C_ROSE; run.bold = True
            else:
                run.font.color.rgb = C_LIGHT
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


# ════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(32)
r_brand = p.add_run('BITSAUTO')
r_brand.bold = True; r_brand.font.size = Pt(42)
r_brand.font.color.rgb = C_CYAN

r_rest = p.add_run('  MONITORING PLATFORM')
r_rest.bold = True; r_rest.font.size = Pt(42)
r_rest.font.color.rgb = C_WHITE

p2 = doc.add_paragraph()
r2 = p2.add_run('Department-Categorised Implementation Status & Roadmap Report')
r2.bold = True; r2.font.size = Pt(18); r2.font.color.rgb = C_GOLD

p3 = doc.add_paragraph()
r3 = p3.add_run('Volume 1 Complete  ·  Volume 2 In Progress  ·  Volume 3 Proposals  ·  May 2026')
r3.font.size = Pt(11); r3.font.color.rgb = C_CYAN_SOFT

p4 = doc.add_paragraph()
r4 = p4.add_run('NOC  ·  Finance & Billing  ·  Commercial  ·  Fraud & Security  ·  Analytics  ·  Engineering')
r4.font.size = Pt(10); r4.font.color.rgb = C_MID

add_rule(doc, C_CYAN)

for k, v in [
    ('Report Date',    '15 May 2026'),
    ('Platform',       'Sippy Softswitch / VoIP NOC Dashboard'),
    ('App URL',        'https://vo-ip-watcher--junaid70.replit.app'),
    ('Stack',          'React 18 + Vite + Express + PostgreSQL (Drizzle ORM) + Replit Auth OIDC'),
    ('Prepared by',    'BitsAuto Engineering'),
    ('Classification', 'Internal — Confidential'),
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f'{k}:  '); r1.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = C_CYAN_SOFT
    r2 = p.add_run(v); r2.font.size = Pt(10); r2.font.color.rgb = C_SLATE

add_rule(doc, C_CYAN)
doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# SECTION 1 — EXECUTIVE SUMMARY
# ════════════════════════════════════════════════════════════════════════════
add_section_banner(doc, '1.  EXECUTIVE SUMMARY', 'Platform overview and delivery status')
add_body(doc, (
    'The BitsAuto VoIP Monitoring Platform is a production-grade NOC dashboard built for Sippy '
    'Softswitch deployments. It delivers real-time call-quality monitoring, telecom KPI tracking, '
    'fraud detection, revenue analytics, team management, and end-to-end client provisioning — '
    'all inside a single role-gated web application with dark/light mode, a command palette, '
    'and a mobile-responsive layout.'
), 10)
add_body(doc, (
    'Volume 1 (24 features across 5 tiers) is 100 % complete. '
    'Volume 2 (24 extended features) has 9 fully live, 5 partially implemented, and 10 not yet started. '
    'This report organises all features by operational department for cross-team clarity, '
    'and proposes 20 additional Volume 3 features.'
), 10)

make_table(doc,
    ['Scope', 'Total Features', 'Live', 'Partial', 'Not Yet Built'],
    [
        ['Volume 1 — Core Platform',               '24', '24  (100 %)', '0',         '0'],
        ['Volume 2 — Extended Features',           '24', '9   (38 %)',  '5  (21 %)', '10  (42 %)'],
        ['Volume 3 — New Proposals (this report)', '20', '—',           '—',         '20  (0 % started)'],
        ['TOTAL',                                  '68', '33',          '5',         '30'],
    ],
    col_widths=[6.0, 3.0, 3.2, 3.0, 3.2]
)

add_heading(doc, 'Feature Count by Department', 2, C_CYAN, size=11, space_before=12)
make_table(doc,
    ['Department', 'Vol 1 (Live)', 'Vol 2 Live', 'Vol 2 Partial', 'Vol 2 Not Built', 'Vol 3 Proposals'],
    [
        ['NOC & Operations',                  '6', '0', '2', '0', '1'],
        ['Finance & Billing',                 '4', '3', '0', '0', '4'],
        ['Commercial & Client Management',    '2', '4', '0', '0', '3'],
        ['Fraud & Security',                  '2', '1', '1', '3', '3'],
        ['Analytics & Business Intelligence', '2', '1', '1', '4', '4'],
        ['Engineering & Infrastructure',      '7', '0', '1', '3', '5'],
        ['UX & Platform',                     '1', '0', '0', '0', '0'],
    ],
    col_widths=[5.2, 2.5, 2.5, 2.5, 3.0, 3.3]
)
doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# SECTION 2 — DEPARTMENT-BY-DEPARTMENT FEATURE OVERVIEW
# ════════════════════════════════════════════════════════════════════════════
add_section_banner(doc, '2.  ALL FEATURES BY DEPARTMENT', 'Status key:  ✅ LIVE  |  ⚠️ PARTIAL  |  ❌ NOT BUILT')
add_body(doc, (
    'Every Vol 1 and Vol 2 feature is grouped below by the operational team that uses or benefits from it most.'
), 10)

dept_features = [
    ('NOC & Operations', [
        ['Vol 1 #1',  'Live Call Quality Monitoring — MOS, Jitter, Latency, Packet Loss', '✅ LIVE',    'Real-time WebSocket metrics; threshold alerts; per-call history'],
        ['Vol 1 #5',  'Alert Engine — threshold-based triggers + email delivery',          '✅ LIVE',    'Gmail SMTP; per-rule cooldowns; alert_rules table; enable/disable'],
        ['Vol 1 #7',  'Traffic Drop Detector — >50 % drop alert',                         '✅ LIVE',    'Background job every 5 min; 30-min cooldown; open/resolved states'],
        ['Vol 1 #8',  'Client Traffic Pulse — per-client live concurrency cards',          '✅ LIVE',    'Live call count + % of peak bar + trend indicator on Graphs page'],
        ['Vol 1 #14', 'Server Monitoring — 6-tab infrastructure view',                    '✅ LIVE',    'TCP probe 60s; outage_log; RTP bandwidth; disk & memory; SIP storm'],
        ['Vol 1 #25', 'Mobile-Responsive NOC View',                                       '✅ LIVE',    'Hamburger Sheet sidebar; push notification opt-in in Settings'],
        ['Vol 2 #1',  'Jitter & Packet Loss Monitor (extended)',                           '⚠️ PARTIAL', '/rtp-analytics exists. Missing: per-carrier trend chart, jitter col in Live Calls'],
        ['Vol 2 #4',  'RTP Port & NAT Traversal Monitor',                                  '⚠️ PARTIAL', '/rtp-analytics partial. Missing: one-way-audio by subnet, STUN/TURN hints'],
    ]),
    ('Finance & Billing', [
        ['Vol 1 #10', 'Revenue & Margin Analytics — 30-day P&L',                          '✅ LIVE',    'Admin: 4-card Revenue/Cost/Profit/Margin. KAM: filtered to own clients'],
        ['Vol 1 #15', 'Vendor Balance Tracker — delta-based cost',                        '✅ LIVE',    '2-hour rolling balance snapshots; cost = balance decrease over window'],
        ['Vol 1 #16', 'Rate Cards & Tariff Management',                                   '✅ LIVE',    'Sippy tariff/rate CRUD; bulk import CSV/XLSX'],
        ['Vol 1 #19', 'Payments & Billing — full Sippy billing API',                      '✅ LIVE',    'Full coverage of Sippy billing API articles 107440–107446'],
        ['Vol 2 #5',  'LCR Analyser',                                                     '✅ LIVE',    '/lcr-analyser — CDR vs vendor rate cards; missed savings in USD'],
        ['Vol 2 #7',  'Vendor SLA Scorecard',                                             '✅ LIVE',    '/vendor-sla-scorecard — per-vendor ASR/ACD/PDD delivery metrics'],
        ['Vol 2 #8',  'Cost Optimisation Engine',                                         '✅ LIVE',    '/cost-optimisation — margin-per-route analysis; worst routes highlighted'],
    ]),
    ('Commercial & Client Management', [
        ['Vol 1 #6',  'KAM Management & Portfolio Assignment',                            '✅ LIVE',    'kams + kam_accounts tables; live call overlay; per-KAM BitsEye drill-down'],
        ['Vol 1 #9',  'Traffic Map — Leaflet choropleth by destination country',          '✅ LIVE',    'CDR country; TopoJSON; violet scale; time-range selector'],
        ['Vol 2 #13', 'Reseller White-Label Portal',                                      '✅ LIVE',    '/reseller — scoped P&L, CDR reports, invoices per reseller account'],
        ['Vol 2 #14', 'SLA Monitoring & Breach Alerting',                                 '✅ LIVE',    '/sla-breaches — per-client SLA compliance; auto breach email to KAM'],
        ['Vol 2 #15', 'Number Inventory Management',                                      '✅ LIVE',    '/number-intelligence — range tracking, utilisation %, expiry, porting'],
        ['Vol 2 #16', 'Click-to-Call & Test Call Launcher',                               '✅ LIVE',    '/test-call — XML-RPC originate; real-time progress; CDR auto-linked'],
    ]),
    ('Fraud & Security', [
        ['Vol 1 #4',  'Role-Based Access Control — Admin / Management / Viewer',          '✅ LIVE',    'Replit OIDC + requireRole() guard on every protected route'],
        ['Vol 1 #12', 'Security Hardening — Helmet, rate-limit, CSP',                    '✅ LIVE',    '300/15min general; 20/15min auth; suspicious IP log (15+ 401/403 in 5 min)'],
        ['Vol 1 #13', 'Fraud / FAS + IRSF Detection Engine',                              '✅ LIVE',    'Zero-billed, high-PDD, short-billed, early-answer; IRSF premium prefix scan'],
        ['Vol 2 #10', 'SIP Brute-Force Login Monitor',                                    '⚠️ PARTIAL', '/firewall exists. Missing: auto-block on >50 failed regs/5 min, geo-origin map'],
        ['Vol 2 #9',  'Robocall & Neighbour Spoofing Detector',                           '❌ NOT BUILT','No sequential CLI detection; no <8 s duration flagging; no auto-block'],
        ['Vol 2 #11', 'Concurrent Call Limit Enforcer',                                   '❌ NOT BUILT','No per-account poll vs limit; no auto-disconnect of excess calls'],
        ['Vol 2 #12', 'Fraud Velocity Scoring',                                           '❌ NOT BUILT','No multi-signal risk score (0–100); no score history chart in NOC sidebar'],
    ]),
    ('Analytics & Business Intelligence', [
        ['Vol 1 #2',  'Call Session Tracking — CDR: ASR, ACD, PDD, CPS',                 '✅ LIVE',    'Sippy XML-RPC getAccountCDRs / getCustomerCDRs; CPS from 1-hr window'],
        ['Vol 2 #6',  'Call Flow Simulator',                                              '✅ LIVE',    '/call-flow-simulator — step-by-step dialplan trace without a real call'],
        ['Vol 2 #19', 'Peer Carrier Benchmarking',                                        '⚠️ PARTIAL', '/carrier-scoring exists. Missing: industry-median comparison, traffic-light vs peers'],
        ['Vol 2 #17', 'Traffic Forecasting Dashboard',                                    '❌ NOT BUILT','No 8-week same-day model; no hourly forecast; no confidence bands or deviation alert'],
        ['Vol 2 #18', 'Geographic Revenue Heatmap',                                       '❌ NOT BUILT','No Revenue/Cost/Margin/Volume world map. Distinct from CDR traffic map'],
        ['Vol 2 #20', 'CDR Anomaly Detector (Statistical)',                               '❌ NOT BUILT','No nightly 3σ deviation job; no anomaly inbox; no severity score per batch'],
        ['Vol 2 #2',  'Codec Negotiation Analytics',                                      '❌ NOT BUILT','No codec chart; no transcoding rate. Requires Sippy CDR codec field — verify first'],
    ]),
    ('Engineering & Infrastructure', [
        ['Vol 1 #3',  'IP Endpoint Probe — TCP SIP port liveness',                        '✅ LIVE',    'Live probe on Server Monitoring; per-host outage_log'],
        ['Vol 1 #11', 'Sippy Change Watcher — IPs, clients, vendors',                    '✅ LIVE',    'Polls 5 min; sippy_snapshots for false-positive prevention; email per event'],
        ['Vol 1 #17', 'DID Management',                                                   '✅ LIVE',    'List, assign, release DIDs via Sippy API'],
        ['Vol 1 #18', 'Multi-Switch Support',                                             '✅ LIVE',    'SwitchesPanel in Settings; push rate changes to all enabled switches'],
        ['Vol 1 #20', 'Trunk & Trunk Connection Management',                              '✅ LIVE',    'Full CRUD for trunks and connections via Sippy API'],
        ['Vol 1 #23', 'API Key Management',                                               '✅ LIVE',    'vw_ prefix; SHA-256 hash; Bearer auth on /ext/ routes; instant revoke'],
        ['Vol 2 #24', 'Multi-Switch Consolidated View',                                   '✅ LIVE',    '/multi-switch — unified dashboard across all configured Sippy instances'],
        ['Vol 2 #22', 'Config Change Tracker (extended)',                                  '⚠️ PARTIAL', 'Sippy Change Watcher live. Missing: diff viewer, rollback suggestion per type'],
        ['Vol 2 #21', 'SIP Registration Database & Viewer',                               '❌ NOT BUILT','No searchable live reg DB; no device fingerprinting; no geography flag'],
        ['Vol 2 #23', 'Maintenance Window Scheduler',                                     '❌ NOT BUILT','No alert suppression for planned downtime; no cron-style schedule UI'],
        ['Vol 2 #3',  'DTMF Analytics & Failure Tracking',                                '❌ NOT BUILT','No DTMF method tracking; no failure-rate per carrier; no cause-code flagging'],
    ]),
    ('UX & Platform', [
        ['Vol 1 #21', 'Dark / Light Mode Toggle',                                         '✅ LIVE',    'CSS vars; ThemeProvider; localStorage persist; default dark on first load'],
        ['Vol 1 #22', 'Quick Actions Command Bar — Ctrl+K / Cmd+K',                       '✅ LIVE',    'shadcn CommandDialog; role-filtered; all nav items searchable'],
        ['Vol 1 #24', 'Customizable Dashboard Widgets',                                   '✅ LIVE',    '5 toggleable sections; prefs per-user in DB; badge shows hidden count'],
    ]),
]

for dept_name, features in dept_features:
    colour, bg = DEPT_COLORS[dept_name]
    add_dept_banner(doc, dept_name)
    make_table(doc,
        ['Ref', 'Feature', 'Status', 'Notes'],
        features,
        col_widths=[1.8, 5.5, 2.2, 6.5],
        header_bg=bg
    )

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# SECTION 3 — 14-DAY ACTIVITY LOG
# ════════════════════════════════════════════════════════════════════════════
add_section_banner(doc, '3.  DEVELOPMENT ACTIVITY LOG — 1 MAY TO 15 MAY 2026', 'Chronological record of all significant work completed')

activity = [
    ('15 May 2026', 'Client Wizard — Complete 5-Step UX Overhaul', 'Commercial & Client Management', [
        'Step 1 (Account Details): Department field replaced with 4 clickable toggle pills. '
        'Company dropdown filtered by selected department. "+ Create Company" inline link added. '
        'Display Name auto-fills from company selection. Username field gains lock/unlock icon. '
        'Password field gains copy-to-clipboard icon. Notification email auto-fills. '
        'A2Z Notification switched from checkbox to toggle.',
        'Step 2 (Billing & Rates): Invoice Template redesigned as 4 clickable tile cards. '
        'Three rate-format dropdowns replaced with multi-select tag pill row (Select All / Clear). '
        'Dialcode Format and Prefix Style replaced with 2-way segmented controls. '
        'Live prefix example updates in real time as options change.',
        'Step 3 (Trunks): Trunk Name replaced with 4 product-class tiles — First Class (amber), '
        'Business Class (blue), Special Charlie (violet), Special Bravo (emerald), plus Custom fallback. '
        'Max Sessions / CPS / Time gain quick-pick chips. Codec selector redesigned as visual tiles. '
        'Media Relay becomes 3-way segmented control. Prefix + CLD live translation preview shown.',
        'Step 4 (IPs): Paste-multiple IP detection — newline / comma / semicolon splits into rows. '
        'Trunk column in IP table is now a dropdown from Step 3 trunk names. '
        'Products section auto-derives from trunk names (read-only) or shows manual tile picker.',
        'Step 5 (Review): Every review section is clickable to jump back to that step. '
        'Live CLD translation preview shown per trunk. Save Draft button on all steps.',
        'Cross-cutting: Enter key advances step. Inline validation fires on blur. '
        'Step breadcrumb pills show a green tick when fully valid.',
    ]),
    ('14 May 2026', 'Company List — KAM Filter Pills + CompanyInfoDialog', 'Commercial & Client Management', [
        'Added KAM quick-filter pills at top of company list, sourced live from /api/kam.',
        'Added CompanyInfoDialog popup — clicking any company name shows: KAM assignment, '
        'location, currency, registered IPs, prefix, products, and trunk config.',
        'Added "New Client Wizard" launch button in Companies page header.',
        'Added product badge chips rendered directly on company list cards.',
    ]),
    ('13–14 May 2026', 'Sippy Username Policy Correction', 'Engineering & Infrastructure', [
        'Removed .toLowerCase() enforcement on Sippy usernames. Sippy confirmed it accepts mixed-case '
        '(e.g. Internal-PTCL provisioned successfully). Username strips only spaces and '
        'invalid characters via .replace(/[^a-zA-Z0-9._-]/g, "") — no forced lowercase.',
    ]),
    ('12–13 May 2026', 'Client Wizard — KAM Dropdown & Routing Group Wiring', 'Commercial & Client Management', [
        'KAM dropdown in company-create.tsx migrated from /api/users to dedicated /api/kam endpoint.',
        'Routing group selector changed from fixed 12-item dropdown to open <input> + <datalist> '
        'combo, allowing any arbitrary RG ID rather than only cached list items.',
    ]),
    ('9–11 May 2026', 'Analytics, Reporting & Routing Pages', 'Analytics & Business Intelligence', [
        'BitsEye per-entity drill-down: per-KAM and per-client concurrency overlays added.',
        'Revenue & Margin Analytics: Admin 4-card P&L view; KAM-filtered revenue for Management role.',
        'QoS Heatmap page (/qos-heatmap): route quality heat visualization by destination and time.',
        'SLA Breaches page (/sla-breaches): per-client SLA compliance tracking with breach history.',
        'Vendor SLA Scorecard page (/vendor-sla-scorecard): per-vendor delivery metrics.',
        'LCR Analyser page (/lcr-analyser): CDR vs rate card cross-reference, missed-savings highlight.',
        'Call Flow Simulator (/call-flow-simulator): dialplan trace tool without placing a real call.',
        'Routing Manager (/routing-manager): route group CRUD + destination set management.',
        'Multi-Switch page (/multi-switch): unified view across all configured Sippy instances.',
    ]),
    ('5–8 May 2026', 'Fraud, Security & Compliance Pages', 'Fraud & Security', [
        'FAS/IRSF Detection Engine: zero-billed, high-PDD, short-billed, early-answer categories + '
        'IRSF prefix scan against curated high-risk prefix list.',
        'Auto-Blacklist system (/firewall): rule-based account/prefix blocking via Sippy XML-RPC.',
        'SIP OPTIONS Monitor: trunk liveness detection via SIP keepalive checks.',
        'WhatsApp Alerts page (/whatsapp-alerts): Telegram/WhatsApp channel support wired to Watcher Recipients.',
        'STIR/SHAKEN compliance dashboard (/stir-shaken): attestation level (A/B/C) tracking and trending.',
        'SIP Trace page (/sip-trace): ladder diagram viewer for full SIP dialog (INVITE → 200 OK → BYE).',
        'Compliance page (/compliance): CALEA/GDPR data handling dashboard.',
    ]),
    ('1–4 May 2026', 'Infrastructure, Auth & UX Hardening', 'Engineering & Infrastructure', [
        'Sippy Change Watcher hardened: sippy_snapshots table prevents false-positive alerts.',
        'Replit Auth OIDC: requireRole() guard applied consistently — isAuthenticated() no longer used directly.',
        'Security hardening: Helmet.js CSP headers tightened; suspicious IP tracker (>15 failed attempts/5 min).',
        'Background job intervals staggered and mutex-guarded — reduces Sippy XML-RPC load by ~65–70 %.',
        'Client Wizard draft save: partial wizard state persisted to DB, resumable on reload.',
        'Command Bar (Ctrl+K): role-filtered navigation added for all new pages added this sprint.',
    ]),
]

for date, title, dept, bullets in activity:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(f'{date}  —  {title}')
    r1.bold = True; r1.font.size = Pt(11); r1.font.color.rgb = C_CYAN
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0); p2.paragraph_format.space_after = Pt(3)
    dept_colour = DEPT_COLORS.get(dept, (C_MID, BG_NAVY_MID))[0]
    r2 = p2.add_run(f'  [ {dept} ]')
    r2.italic = True; r2.font.size = Pt(9); r2.font.color.rgb = dept_colour
    for b in bullets:
        add_bullet(doc, b, 9.5, C_SLATE)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# SECTION 4 — VOLUME 1 STATUS
# ════════════════════════════════════════════════════════════════════════════
add_section_banner(doc, '4.  VOLUME 1 — FULL FEATURE STATUS  (All 24 Features, by Department)', 'All 24 features across Tiers 1–5 are confirmed live as of 15 May 2026')

v1_rows = [
    ['T1', '1',  'NOC & Operations',                 'Live Call Quality Monitoring — MOS, Jitter, Latency, Packet Loss', 'LIVE', 'Real-time WebSocket metrics, threshold alerts, per-call history'],
    ['T1', '2',  'Analytics',                        'Call Session Tracking — CDR: ASR, ACD, PDD, CPS',                 'LIVE', 'Sippy XML-RPC getAccountCDRs / getCustomerCDRs; CPS from 1-hour window'],
    ['T1', '3',  'Engineering',                      'IP Endpoint Probe — TCP SIP port liveness check',                  'LIVE', 'Live probe status on Server Monitoring; per-host outage_log'],
    ['T1', '4',  'Fraud & Security',                 'Role-Based Access Control — Admin / Management / Viewer',           'LIVE', 'Replit OIDC + requireRole() guard on every protected route'],
    ['T1', '5',  'NOC & Operations',                 'Alert Engine — threshold-based triggers + email delivery',          'LIVE', 'Gmail SMTP; per-rule cooldowns; alert_rules table with enable/disable'],
    ['T2', '6',  'Commercial & Client Management',   'KAM Management & Portfolio Assignment',                             'LIVE', 'kams + kam_accounts tables; live call overlay; BitsEye per-KAM drill-down'],
    ['T2', '7',  'NOC & Operations',                 'Traffic Drop Detector — >50 % drop alert',                          'LIVE', 'Background job every 5 min; 30-min cooldown; open/resolved states'],
    ['T2', '8',  'NOC & Operations',                 'Client Traffic Pulse — per-client live concurrency cards',          'LIVE', 'Live call count + % of peak bar + trend indicator on Graphs page'],
    ['T2', '9',  'Commercial & Client Management',   'Traffic Map — Leaflet choropleth by destination country',           'LIVE', 'CDR country field; TopoJSON; violet scale; time-range selector'],
    ['T2', '10', 'Finance & Billing',                'Revenue & Margin Analytics — 30-day P&L',                           'LIVE', 'Admin: 4-card Revenue/Cost/Profit/Margin. KAM: filtered to own clients'],
    ['T3', '11', 'Engineering',                      'Sippy Change Watcher — IP rules, clients, vendors',                 'LIVE', 'Polls 5 min; sippy_snapshots false-positive prevention; email per event'],
    ['T3', '12', 'Fraud & Security',                 'Security Hardening — Helmet, rate-limit, CSP',                      'LIVE', '300/15min general; 20/15min auth; suspicious IP log (15+ 401/403 in 5 min)'],
    ['T3', '13', 'Fraud & Security',                 'Fraud / FAS + IRSF Detection',                                      'LIVE', 'Zero-billed, high-PDD, short-billed, early-answer; IRSF premium prefix scan'],
    ['T3', '14', 'Engineering',                      'Server Monitoring — 6-tab infrastructure view',                     'LIVE', 'TCP probe 60s; outage_log; RTP bandwidth; disk & memory; SIP reg storm'],
    ['T3', '15', 'Finance & Billing',                'Vendor Balance Tracker — delta-based cost',                         'LIVE', '2-hour rolling balance snapshots; cost = balance decrease over window'],
    ['T4', '16', 'Finance & Billing',                'Rate Cards & Tariff Management',                                    'LIVE', 'Sippy tariff/rate CRUD; bulk import from CSV/XLSX files'],
    ['T4', '17', 'Engineering',                      'DID Management',                                                    'LIVE', 'List, assign, release DIDs via Sippy API'],
    ['T4', '18', 'Engineering',                      'Multi-Switch Support',                                              'LIVE', 'SwitchesPanel in Settings; push rate changes to all enabled switches'],
    ['T4', '19', 'Finance & Billing',                'Payments & Billing',                                                'LIVE', 'Full Sippy billing API coverage (articles 107440–107446)'],
    ['T4', '20', 'Engineering',                      'Trunk & Trunk Connection Management',                               'LIVE', 'Full CRUD for trunks and connections via Sippy API'],
    ['T5', '21', 'UX & Platform',                   'Dark / Light Mode Toggle',                                          'LIVE', 'CSS vars; ThemeProvider; localStorage persist; default dark on first load'],
    ['T5', '22', 'UX & Platform',                   'Quick Actions Command Bar — Ctrl+K / Cmd+K',                        'LIVE', 'shadcn CommandDialog; role-filtered; all nav items searchable'],
    ['T5', '23', 'Engineering',                      'API Key Management',                                                'LIVE', 'vw_ prefix; SHA-256 hash; Bearer auth on /ext/ routes; revoke instantly'],
    ['T5', '24', 'UX & Platform',                   'Customizable Dashboard Widgets',                                    'LIVE', '5 toggleable sections; prefs per-user in DB; badge shows hidden count'],
    ['T5', '25', 'NOC & Operations',                 'Mobile-Responsive NOC View',                                        'LIVE', 'Hamburger Sheet sidebar; push notification opt-in in Settings'],
]
make_table(doc,
    ['Tier', '#', 'Department', 'Feature', 'Status', 'Implementation Notes'],
    v1_rows,
    col_widths=[0.9, 0.7, 3.2, 5.0, 1.7, 5.5]
)
doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# SECTION 5 — VOLUME 2 STATUS
# ════════════════════════════════════════════════════════════════════════════
add_section_banner(doc, '5.  VOLUME 2 — EXTENDED FEATURE STATUS  (All 24 Features, by Department)',
                   '9 fully live · 5 partially implemented · 10 not yet started')

v2_rows = [
    ['NOC & Operations',                  '1',  'Jitter & Packet Loss Monitor',          '⚠️ PARTIAL',    '/rtp-analytics exists. Missing: per-carrier trend chart, jitter col in Live Calls'],
    ['NOC & Operations',                  '4',  'RTP Port & NAT Traversal Monitor',       '⚠️ PARTIAL',    '/rtp-analytics partial. Missing: one-way-audio by subnet, STUN/TURN hints'],
    ['Finance & Billing',                 '5',  'LCR Analyser',                           '✅ LIVE',       '/lcr-analyser — CDR vs all active vendor rate cards; missed savings in USD'],
    ['Finance & Billing',                 '6',  'Call Flow Simulator',                    '✅ LIVE',       '/call-flow-simulator — step-by-step dialplan trace without placing a real call'],
    ['Finance & Billing',                 '7',  'Vendor SLA Scorecard',                   '✅ LIVE',       '/vendor-sla-scorecard — per-vendor ASR/ACD/PDD delivery metrics'],
    ['Finance & Billing',                 '8',  'Cost Optimisation Engine',               '✅ LIVE',       '/cost-optimisation — margin-per-route analysis; worst routes highlighted'],
    ['Fraud & Security',                  '9',  'Robocall & Neighbour Spoofing Detector', '❌ NOT BUILT',  'No sequential CLI detection; no <8 s duration flagging; no auto-block via XML-RPC'],
    ['Fraud & Security',                  '10', 'SIP Brute-Force Login Monitor',          '⚠️ PARTIAL',    '/firewall exists. Missing: auto-add-to-blocklist on >50 failed regs/5 min'],
    ['Fraud & Security',                  '11', 'Concurrent Call Limit Enforcer',         '❌ NOT BUILT',  'No per-account poll vs configured limit; no auto-disconnect of excess calls'],
    ['Fraud & Security',                  '12', 'Fraud Velocity Scoring',                 '❌ NOT BUILT',  'No multi-signal risk score (0–100); no score history chart in NOC sidebar'],
    ['Commercial & Client Management',    '13', 'Reseller White-Label Portal',            '✅ LIVE',       '/reseller — scoped P&L, CDR reports, and invoices per reseller account'],
    ['Commercial & Client Management',    '14', 'SLA Monitoring & Breach Alerting',       '✅ LIVE',       '/sla-breaches — per-client SLA compliance; auto breach email to KAM and client'],
    ['Commercial & Client Management',    '15', 'Number Inventory Management',            '✅ LIVE',       '/number-intelligence — range tracking, utilisation %, expiry, porting tracker'],
    ['Commercial & Client Management',    '16', 'Click-to-Call & Test Call Launcher',     '✅ LIVE',       '/test-call — XML-RPC originate; real-time call progress; CDR auto-linked'],
    ['Analytics',                         '17', 'Traffic Forecasting Dashboard',          '❌ NOT BUILT',  'No 8-week same-day model; no hourly forecast; no confidence bands; no deviation alert'],
    ['Analytics',                         '18', 'Geographic Revenue Heatmap',             '❌ NOT BUILT',  'No Revenue/Cost/Margin/Volume world map. Distinct from existing CDR traffic map'],
    ['Analytics',                         '19', 'Peer Carrier Benchmarking',              '⚠️ PARTIAL',    '/carrier-scoring exists. Missing: industry-median comparison, traffic-light vs peers'],
    ['Analytics',                         '20', 'CDR Anomaly Detector (Statistical)',     '❌ NOT BUILT',  'No nightly 3σ deviation job; no anomaly inbox in Reports; no severity score per batch'],
    ['Analytics',                         '2',  'Codec Negotiation Analytics',            '❌ NOT BUILT',  'No codec usage chart; no transcoding-rate-per-carrier. Needs Sippy CDR codec field'],
    ['Engineering & Infrastructure',      '21', 'SIP Registration Database & Viewer',     '❌ NOT BUILT',  'No searchable live reg DB; no device fingerprinting; no unexpected-geography flag'],
    ['Engineering & Infrastructure',      '22', 'Config Change Tracker (extended)',        '⚠️ PARTIAL',    'Sippy Change Watcher live. Missing: diff viewer, rollback suggestion per config type'],
    ['Engineering & Infrastructure',      '23', 'Maintenance Window Scheduler',            '❌ NOT BUILT',  'No alert suppression for planned downtime; no cron-style schedule UI'],
    ['Engineering & Infrastructure',      '24', 'Multi-Switch Consolidated View',          '✅ LIVE',       '/multi-switch — unified dashboard across all configured Sippy instances'],
    ['Engineering & Infrastructure',      '3',  'DTMF Analytics & Failure Tracking',       '❌ NOT BUILT',  'No DTMF method tracking; no failure-rate-per-carrier; no cause-code flagging'],
]
make_table(doc,
    ['Department', '#', 'Feature', 'Status', 'Gap / Implementation Notes'],
    v2_rows,
    col_widths=[3.2, 0.7, 4.5, 2.2, 5.4]
)
doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# SECTION 6 — FEATURES REMAINING
# ════════════════════════════════════════════════════════════════════════════
add_section_banner(doc, '6.  FEATURES REMAINING TO COMPLETE — BY DEPARTMENT',
                   '15 items either unstarted or with meaningful gaps — listed in recommended build order')

remaining_by_dept = [
    ('NOC & Operations', C_CYAN, BG_NAVY_MID, [
        ('MOS Score Trending & Quality Reports',
         'Vol 1 #4 (gap)  |  Effort: Medium  |  Value: High',
         'The platform shows real-time MOS per call but lacks historical analysis. '
         'Build: MOS trend chart over 24h/7d/30d broken down by carrier and destination. '
         '"Quality Events" log: any 15-minute window where average MOS drops below 3.5. '
         'Carrier Quality League Table: rank vendors by average MOS over selected period. '
         'Auto-alert when a carrier\'s rolling 1-hour MOS drops below configurable threshold.'),
        ('Traffic Anomaly Detector — Statistical Baseline',
         'Vol 1 #7 (gap)  |  Effort: Medium  |  Value: High',
         'Existing Traffic Drop Detector only fires at >50 % drop thresholds. '
         'Extend with 14-day per-hour rolling baseline model. Alert when CPS or concurrent calls '
         'exceed 2 standard deviations above the baseline. Distinguish business-hours peak from a 3 AM spike. '
         'Display anomaly timeline overlay on the Graphs page.'),
    ]),
    ('Finance & Billing', C_GREEN, '063B2A', [
        ('Billing Report Generator — PDF / Excel',
         'Vol 1 #12  |  Effort: Medium  |  Value: High',
         'Select client + billing period + rate card → aggregate CDRs by destination, apply rates, produce '
         'a formatted invoice with: summary totals (minutes, cost per destination group) and an itemised '
         'CDR attachment. Export as PDF (with company logo) or Excel. Send directly from within the app.'),
    ]),
    ('Fraud & Security', C_ROSE, '3B0A1A', [
        ('Fraud Velocity Scoring',
         'Vol 2 #12  |  Effort: Medium  |  Value: Critical',
         'Real-time multi-signal risk score (0–100) per account. Score components: destination risk level, '
         'call velocity (calls/minute), CLI diversity (unique caller IDs per hour), time-of-day anomaly, '
         'short-duration ratio (% calls < 8 seconds). '
         'Score displayed live in NOC dashboard sidebar. Accounts above threshold (75) trigger auto-alert. '
         'Score history charted per account on Client detail page.'),
        ('Robocall & Neighbour Spoofing Detector',
         'Vol 2 #9  |  Effort: Medium  |  Value: High',
         'Extends FAS engine with: (1) Sequential CLI detection — flag accounts using caller IDs in '
         'arithmetic sequences (+1-555-0001, +1-555-0002…); '
         '(2) Neighbour spoofing — caller ID prefix matches dialled number\'s area code. '
         'Also flag accounts with >200 calls/hour where average duration < 8 seconds. '
         'Auto-block identified source accounts via Sippy XML-RPC.'),
        ('Concurrent Call Limit Enforcer',
         'Vol 2 #11  |  Effort: Low  |  Value: Medium',
         'Poll live calls per account every 10 seconds against the configured limit stored in Sippy. '
         'Alert NOC when any account exceeds its limit by >20 %. '
         'Optional auto-disconnect: terminate the oldest excess call via XML-RPC. '
         'Limit breach history log per account.'),
    ]),
    ('Analytics & Business Intelligence', C_AMBER, '3B2200', [
        ('CDR Anomaly Detector — Statistical',
         'Vol 2 #20  |  Effort: Medium  |  Value: High',
         'Nightly backend job comparing previous day\'s CDRs against rolling 30-day per-account baseline. '
         'Flags records deviating >3 standard deviations in: call duration, cost per minute, or destination '
         'distribution. Anomaly inbox added to Reports page with severity score per flagged batch.'),
        ('Traffic Forecasting Dashboard',
         'Vol 2 #17  |  Effort: High  |  Value: Medium',
         'Analyse same-day-of-week patterns over past 8 weeks. '
         'Generate an hourly traffic forecast for next 24 hours with confidence bands '
         '(optimistic / expected / pessimistic). Alert if actual traffic deviates >30 % from forecast.'),
        ('Geographic Revenue Heatmap',
         'Vol 2 #18  |  Effort: Medium  |  Value: Medium',
         'World map where colour intensity = revenue generated per country. '
         'Toggle between Revenue / Cost / Margin / Call Volume views. '
         'Click a country to drill into that country\'s CDRs, carrier breakdown, and ASR trend. '
         'Distinct from the existing CDR traffic map — focused on the financial layer.'),
        ('Codec Negotiation Analytics',
         'Vol 2 #2  |  Effort: Medium  |  Value: Medium',
         'Parse codec fields from Sippy CDRs. Display codec usage breakdown (pie/bar chart). '
         'Show transcoding rate per carrier. Flag vendors forcing costly transcode paths. '
         'Note: requires Sippy instance to expose codec field in CDR API — verify before building.'),
    ]),
    ('Engineering & Infrastructure', C_TEAL, '053030', [
        ('SIP Registration Database & Viewer',
         'Vol 2 #21  |  Effort: Low  |  Value: Medium',
         'A searchable real-time list of all currently registered SIP endpoints pulled from Sippy. '
         'Fields: username, IP address, user-agent string, registration expiry timestamp. '
         'Flag registrations from unexpected geographies. Flag unusual user-agent strings (e.g. SIPVicious).'),
        ('Maintenance Window Scheduler',
         'Vol 2 #23  |  Effort: Low  |  Value: Medium',
         'Define recurring or one-off maintenance windows. '
         'All alerts automatically suppressed during the window. '
         'NOC dashboard shows a "Maintenance Mode" banner with countdown during active window.'),
        ('Two-Factor Authentication — TOTP',
         'Vol 1 #19  |  Effort: Medium  |  Value: High',
         'TOTP-based 2FA compatible with Google Authenticator and Authy. '
         'QR code enrollment flow on first login after 2FA is enabled. '
         'Recovery codes generated at enrollment. Admin can enforce 2FA for all Admin/Management roles.'),
    ]),
    ('UX & Platform', C_MID, '1A2535', [
        ('Scheduled Report Delivery',
         'Vol 1 #13  |  Effort: Low  |  Value: High',
         'Automated cron-style delivery of daily/weekly/monthly reports. '
         'Builds on existing infrastructure: Watcher Recipients (live), Gmail SMTP (live), '
         'report-generation logic (already in /reports page). '
         'Deliverables: daily traffic summary at 06:00; weekly ASR/ACD/MOS quality report; '
         'monthly billing summary per client; FAS/fraud incident digest.'),
        ('Audit Log Viewer',
         'Vol 1 #18  |  Effort: Low  |  Value: High',
         'Searchable, tamper-evident log of all administrative actions in the platform. '
         'Events to log: user logins and failed logins; config changes; manual call disconnections; '
         'account blocks/unblocks; alert rule changes; watcher recipient changes. '
         'New DB table: audit_log (id, userId, action, target, before, after, ipAddress, timestamp). '
         'Display: filterable table by user, action type, and date range. Exportable to CSV.'),
    ]),
]

for dept_name, colour, bg, items in remaining_by_dept:
    add_dept_banner(doc, dept_name)
    for title, meta, detail in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(0)
        r = p.add_run(f'  {title}')
        r.bold = True; r.font.size = Pt(10); r.font.color.rgb = C_WHITE
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(0); p2.paragraph_format.space_after = Pt(1)
        p2.paragraph_format.left_indent = Cm(0.8)
        r2 = p2.add_run(meta)
        r2.font.size = Pt(9); r2.italic = True; r2.font.color.rgb = colour
        p3 = doc.add_paragraph()
        p3.paragraph_format.space_before = Pt(1); p3.paragraph_format.space_after = Pt(8)
        p3.paragraph_format.left_indent = Cm(0.8)
        r3 = p3.add_run(detail)
        r3.font.size = Pt(9.5); r3.font.color.rgb = C_SLATE

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# SECTION 7 — VOLUME 3: NEW FEATURE PROPOSALS
# ════════════════════════════════════════════════════════════════════════════
add_section_banner(doc, '7.  VOLUME 3 — NEW FEATURE PROPOSALS BY DEPARTMENT  (20 New Features)',
                   'First-time proposals addressing operational gaps, enterprise requirements, and competitive parity')

v3_sections = [
    ('NOC & Operations', C_CYAN, BG_NAVY_MID, [
        ('AI-Powered NOC Assistant (LLM Chat)',
         'Critical', 'High',
         'A conversational interface NOC engineers query in plain English: '
         '"Show me accounts with calls to >10 new destinations this week", '
         '"Which carriers degraded most in the last 48 hours?", "Summarise all open fraud alerts." '
         'Powered by an LLM with read access to the platform data API. '
         'Provides plain-language incident summaries and suggests remediation actions.'),
    ]),
    ('Finance & Billing', C_GREEN, '063B2A', [
        ('Automated Invoice Delivery & Payment Collection',
         'High', 'High',
         'Generate PDF invoices at billing cycle end and email them to the client automatically. '
         'Include a hosted payment link. Track payment status: paid / overdue / disputed. '
         'Automatic balance lock when an account reaches the overdue threshold.'),
        ('Carrier Invoice Reconciliation',
         'Critical', 'High',
         'Upload a carrier invoice (PDF/CSV). Platform extracts line items and cross-references '
         'against BitsAuto CDRs for the same period. Flags: overbilled minutes, wrong rates, missing CDRs. '
         'Generates a formatted dispute report pre-populated with all discrepancies.'),
        ('Profit & Loss Forecasting',
         'High', 'Medium',
         'Projects next-month P&L using: current traffic growth trend, contracted sell rates, and '
         'current carrier buy rates. Presents rate-renegotiation impact scenarios. '
         'Supports KAM and commercial team preparation for quarterly business reviews.'),
        ('Monthly Executive PDF Report',
         'High', 'Medium',
         'Auto-generated one-page PDF delivered to a configurable email list on the 1st of each month. '
         'Content: total minutes, revenue, gross margin, top 10 destinations by volume, '
         'fraud events count, carrier quality league table, and month-over-month trend spark charts.'),
    ]),
    ('Commercial & Client Management', C_CYAN_SOFT, '0D2A4A', [
        ('Client Self-Service Portal — White-Label',
         'Critical', 'High',
         'A fully white-labeled web portal that end-clients log into with their own branding. '
         'Clients see their own CDRs, live call count, balance, and rate cards. '
         'They can top up balance, download invoices, and raise support tickets.'),
        ('Client Onboarding Kanban Tracker',
         'Medium', 'Low',
         'Kanban board showing every client in the provisioning pipeline: '
         'Draft → IPs Submitted → IPs Approved → Provisioned → Test Call Done → Live. '
         'Each card shows responsible KAM, pending blockers, and last-action timestamp. '
         'Linked directly to the Client Wizard draft system.'),
        ('Predictive Client Churn Scoring',
         'High', 'Medium',
         'Uses CDR traffic trends to identify clients whose call volume is declining or who have '
         'had repeated quality incidents. Generates a churn risk score per client (updated daily). '
         'KAM dashboard surfaces "At Risk" clients with one-click outreach suggestion.'),
    ]),
    ('Fraud & Security', C_ROSE, '3B0A1A', [
        ('Zero-Trust Access Model',
         'High', 'High',
         'Every API call validated not just by role but by registered device, IP range, and time-of-day. '
         'Admin actions from unrecognised IPs require secondary approval from another admin. '
         'Session tokens bound to IP; re-authentication required on IP change.'),
        ('Threat Intelligence Feed Integration',
         'High', 'Medium',
         'Subscribe to live threat feeds (e.g. Spamhaus, ThreatCrowd) and cross-reference inbound '
         'calling numbers and source IPs against known-bad lists in real time. '
         'Auto-block IPs flagged as VoIP fraud infrastructure. Daily threat feed refresh with email summary.'),
        ('SOC 2 / ISO 27001 Compliance Checklist',
         'Medium', 'High',
         'Self-assessment checklist mapped to SOC 2 Type II controls. '
         'Platform auto-fills evidence for controls it can verify (audit log present, MFA enforced, '
         'encryption at rest). Exports a compliance readiness PDF for auditors.'),
    ]),
    ('Analytics & Business Intelligence', C_AMBER, '3B2200', [
        ('Automated Route Optimisation Recommendations',
         'High', 'Medium',
         'Analyses current routing priorities, carrier MOS/ASR trends, and rate card margins to '
         'recommend: which carriers to de-prioritise, which routes to renegotiate, and when to add '
         'capacity. Presented as actionable cards with estimated monthly savings and one-click apply.'),
        ('Regulatory Reporting Module — CALEA / GDPR',
         'Medium', 'High',
         'CALEA: lawful intercept report generation in required regulatory format. '
         'GDPR: CDR anonymisation after configurable retention period; right-to-erasure workflow; '
         'data residency controls flagging which CDRs contain EU personal data.'),
        ('Geographic Revenue Heatmap (Vol 3 enhanced)',
         'Medium', 'Medium',
         'Enhanced version of Vol 2 #18 — adds drill-down to per-carrier breakdown per country, '
         'QBR slide export, and revenue-vs-cost toggle with margin overlay.'),
        ('CDR Anomaly Detector (enhanced — Vol 3)',
         'High', 'Medium',
         'Extends Vol 2 #20 with real-time anomaly detection (not just nightly): streaming 15-minute '
         'window analysis with instant NOC alert when any account deviates from its rolling baseline '
         'by more than 3σ. Feeds directly into the AI NOC Assistant context.'),
    ]),
    ('Engineering & Infrastructure', C_TEAL, '053030', [
        ('SBC / Firewall Rule Sync',
         'High', 'Medium',
         'When a new IP is approved in BitsAuto, push the firewall rule automatically via API '
         '(iptables, Cloudflare, or FortiGate). Keeps Sippy ACL and external firewall in sync. '
         'Full audit trail of every rule push with one-click rollback.'),
        ('Configuration Backup & Point-in-Time Restore',
         'High', 'Medium',
         'Daily snapshot of all platform configuration: routing groups, rate cards, alert rules, '
         'account settings, IP rules — stored as versioned JSON. '
         'One-click restore to any previous snapshot. Critical for disaster recovery.'),
        ('Call Recording Management',
         'Medium', 'Medium',
         'Browse, search, and stream call recordings stored on the Sippy server. '
         'Filter by account, date range, destination, or duration. '
         'Download individual recordings or bulk export. '
         'Configurable retention policy with auto-delete after N days.'),
        ('Scheduled Bulk Actions',
         'Medium', 'Low',
         'Schedule one-off or recurring bulk operations with admin approval gates: '
         '"Block all zero-traffic accounts every Sunday 02:00", '
         '"Push updated rate card to all switches every Monday 00:00". '
         'Full action log with before/after state captured.'),
        ('CRM Integration — HubSpot / Salesforce',
         'High', 'Medium',
         'Sync client records between BitsAuto and the operator CRM. '
         'When a client is provisioned, a Company record is created in HubSpot. '
         'CDR revenue data automatically updates the CRM deal value. '
         'KAM assignment in BitsAuto reflects in CRM account owner field.'),
        ('Webhook Push API for External Systems',
         'High', 'Low',
         'Extend the existing API key system with outbound webhook support. '
         'Configure a destination URL per event type: alert fired, fraud detected, balance low, '
         'client provisioned. Signed payloads with HMAC-SHA256. Retry with exponential back-off.'),
    ]),
]

for cat_name, colour, hdr_bg, features in v3_sections:
    add_dept_banner(doc, cat_name)
    make_table(doc,
        ['Feature', 'Business Value', 'Build Effort', 'Description'],
        [[f[0], f[1], f[2], f[3]] for f in features],
        col_widths=[4.2, 2.0, 2.0, 7.8],
        header_bg=hdr_bg
    )

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# SECTION 8 — FULL PRIORITY MATRIX
# ════════════════════════════════════════════════════════════════════════════
add_section_banner(doc, '8.  COMBINED PRIORITY MATRIX — ALL 35 REMAINING & PROPOSED FEATURES',
                   '5 recommended build sprints with department labelling')

priority_rows = [
    ['S1 — Now',  'NOC & Operations',                 'Vol 2 Gap', 'Traffic Anomaly Detector (baseline)',        'High',     'Medium', 'Catch toll fraud earlier than threshold alerts'],
    ['S1 — Now',  'Finance & Billing',                'Vol 2 Gap', 'Billing Report Generator PDF/Excel',         'High',     'Medium', 'Eliminates manual invoice generation'],
    ['S1 — Now',  'Fraud & Security',                 'Vol 2 Gap', 'Fraud Velocity Scoring',                     'Critical', 'Medium', 'Multi-signal risk; closes fraud detection gap'],
    ['S1 — Now',  'Fraud & Security',                 'Vol 2 Gap', 'CDR Anomaly Detector (Statistical)',         'Critical', 'Medium', 'Nightly job; catch-all for unknown fraud patterns'],
    ['S1 — Now',  'Finance & Billing',                'Vol 3 New', 'Carrier Invoice Reconciliation',             'Critical', 'High',   'Direct cost saving; closes billing dispute loop'],
    ['S1 — Now',  'UX & Platform',                   'Vol 2 Gap', 'Scheduled Report Delivery',                  'Critical', 'Low',    'Fast win; entirely on existing email + reports infra'],
    ['S1 — Now',  'UX & Platform',                   'Vol 2 Gap', 'Audit Log Viewer',                           'Critical', 'Low',    'Enterprise security requirement; new DB table only'],
    ['S2 — Next', 'NOC & Operations',                 'Vol 2 Gap', 'MOS Score Trending (historical)',            'High',     'Medium', 'Proactive quality management; carrier QBR evidence'],
    ['S2 — Next', 'Commercial & Client Management',   'Vol 3 New', 'Client Self-Service Portal',                 'Critical', 'High',   'Reduces operator support load; client satisfaction'],
    ['S2 — Next', 'Finance & Billing',                'Vol 3 New', 'Monthly Executive PDF Report',               'High',     'Medium', 'Management reporting without login'],
    ['S2 — Next', 'Finance & Billing',                'Vol 3 New', 'Automated Invoice Delivery & Payments',      'High',     'High',   'Closes billing loop; reduces DSO'],
    ['S3',        'Fraud & Security',                 'Vol 2 Gap', 'Robocall & Neighbour Spoofing Detector',     'High',     'Medium', 'Extends FAS; protects carrier reputation'],
    ['S3',        'Engineering & Infrastructure',     'Vol 2 Gap', 'SIP Registration Database & Viewer',        'Medium',   'Low',    'Engineering tool; fast to build'],
    ['S3',        'Engineering & Infrastructure',     'Vol 3 New', 'SBC / Firewall Rule Sync',                   'High',     'Medium', 'Removes manual IP management step'],
    ['S3',        'Engineering & Infrastructure',     'Vol 3 New', 'Webhook Push API',                           'High',     'Low',    'Ecosystem play; enables Zapier/Make automation'],
    ['S3',        'Engineering & Infrastructure',     'Vol 3 New', 'Configuration Backup & Restore',             'High',     'Medium', 'Disaster recovery; config versioning'],
    ['S4',        'Analytics',                        'Vol 2 Gap', 'Geographic Revenue Heatmap',                'Medium',   'Medium', 'Executive-level financial geography view'],
    ['S4',        'Analytics',                        'Vol 2 Gap', 'Traffic Forecasting Dashboard',              'Medium',   'High',   'Predictive capacity planning'],
    ['S4',        'NOC & Operations',                 'Vol 3 New', 'AI NOC Assistant (LLM)',                     'High',     'High',   'Natural-language querying; incident summaries'],
    ['S4',        'Engineering & Infrastructure',     'Vol 3 New', 'CRM Integration (HubSpot / Salesforce)',     'High',     'Medium', 'Syncs client + revenue data to CRM'],
    ['S4',        'Finance & Billing',                'Vol 3 New', 'P&L Forecasting',                            'High',     'Medium', 'Commercial QBR preparation tool'],
    ['S4',        'Analytics',                        'Vol 3 New', 'Auto Route Optimisation Recommendations',    'High',     'Medium', 'AI-driven cost reduction recommendations'],
    ['S5',        'Fraud & Security',                 'Vol 2 Gap', 'Concurrent Call Limit Enforcer',            'Medium',   'Low',    'Per-account real-time enforcement'],
    ['S5',        'Analytics',                        'Vol 2 Gap', 'Codec Negotiation Analytics',               'Low',      'Medium', 'Needs Sippy CDR codec field — verify first'],
    ['S5',        'Engineering & Infrastructure',     'Vol 2 Gap', 'Maintenance Window Scheduler',               'Low',      'Low',    'Alert suppression for planned downtime'],
    ['S5',        'Engineering & Infrastructure',     'Vol 2 Gap', 'Two-Factor Authentication (TOTP)',           'Medium',   'Medium', 'Security hardening for high-privilege roles'],
    ['S5',        'Fraud & Security',                 'Vol 3 New', 'Zero-Trust Access Model',                    'High',     'High',   'IP-bound sessions; secondary approval for admin'],
    ['S5',        'Fraud & Security',                 'Vol 3 New', 'Threat Intelligence Feed Integration',       'High',     'Medium', 'Automated IP + number blacklisting'],
    ['S5',        'Fraud & Security',                 'Vol 3 New', 'Regulatory Reporting (CALEA / GDPR)',        'Medium',   'High',   'Required for enterprise and government contracts'],
    ['S5',        'Commercial & Client Management',   'Vol 3 New', 'Predictive Client Churn Scoring',            'Medium',   'Medium', 'KAM retention; outreach trigger'],
    ['S5',        'Commercial & Client Management',   'Vol 3 New', 'Client Onboarding Kanban Tracker',           'Medium',   'Low',    'Provisioning pipeline visibility'],
    ['S5',        'Engineering & Infrastructure',     'Vol 3 New', 'Scheduled Bulk Actions',                     'Low',      'Medium', 'Operational automation for routine tasks'],
    ['S5',        'Engineering & Infrastructure',     'Vol 3 New', 'Call Recording Management',                  'Medium',   'Medium', 'CDR → recording drill-down workflow'],
    ['S5',        'Fraud & Security',                 'Vol 3 New', 'SOC 2 / ISO 27001 Checklist',               'Medium',   'High',   'Compliance readiness export for auditors'],
]

make_table(doc,
    ['Sprint', 'Department', 'Source', 'Feature', 'Business Value', 'Effort', 'Rationale'],
    priority_rows,
    col_widths=[1.6, 3.0, 1.8, 4.5, 2.0, 1.6, 4.5]
)
doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# SECTION 9 — ARCHITECTURE & STACK
# ════════════════════════════════════════════════════════════════════════════
add_section_banner(doc, '9.  ARCHITECTURE & STACK REFERENCE')

make_table(doc,
    ['Layer', 'Technology', 'Notes'],
    [
        ['Frontend',    'React 18 + Vite + TailwindCSS v3',  'SPA; shadcn/ui components; Wouter routing; TanStack Query v5'],
        ['Backend',     'Node.js 20 + Express + TypeScript', 'server/routes.ts (~22,000+ lines); staggered background jobs'],
        ['Database',    'PostgreSQL via Drizzle ORM',         'shared/schema.ts; createInsertSchema from drizzle-zod; Neon'],
        ['Auth',        'Replit OpenID Connect (OIDC)',       'Server-side sessions; requireRole() guard on all protected routes'],
        ['Sippy API',   'XML-RPC over HTTPS + HTTP Digest',  'RFC-2617 2-step probe+auth; server/sippy.ts; creds pair rotation'],
        ['WebSockets',  'Express + WS',                       'Push-based NOC view; live calls; cache-first to reduce Sippy load'],
        ['Email',       'Gmail SMTP (Nodemailer)',             'Alert engine; watcher recipients; billing notifications'],
        ['Deployment',  'Replit (containerised)',              'npm run dev starts Express + Vite on same port'],
        ['Security',    'Helmet.js + express-rate-limit',     '300/15min general; 20/15min auth; CSP + XSS protection headers'],
    ],
    col_widths=[2.8, 4.8, 8.4],
    header_bg=BG_NAVY_MID
)

add_heading(doc, 'Key Source Files', 2, C_CYAN, size=11, space_before=12)
for fpath, desc in [
    ('server/routes.ts',                       '~22,000+ lines — all Express routes, Sippy integration, background polling jobs'),
    ('server/sippy.ts',                        'Sippy XML-RPC client; sippyXmlCredsPairs(); credential swap resilience'),
    ('server/storage.ts',                      'IStorage interface + DatabaseStorage implementation (all DB operations)'),
    ('shared/schema.ts',                       'Drizzle schema, Zod insert schemas, TypeScript select types for all tables'),
    ('client/src/pages/dashboard.tsx',         'NOC dashboard — all real-time KPI and quality widgets'),
    ('client/src/components/layout-shell.tsx', 'Sidebar, mobile hamburger nav, command bar (Ctrl+K), theme toggle'),
    ('client/src/pages/client-wizard.tsx',     'Full 5-step client provisioning wizard (major overhaul — May 2026)'),
    ('client/src/pages/company-list.tsx',      'Company list with KAM filter pills + CompanyInfoDialog popup'),
    ('client/src/pages/fraud.tsx',             'FAS + IRSF detection engine; auto-blacklist rule management'),
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f'  {fpath}  ')
    r1.bold = True; r1.font.size = Pt(9.5); r1.font.color.rgb = C_CYAN_SOFT
    r2 = p.add_run(f'— {desc}')
    r2.font.size = Pt(9.5); r2.font.color.rgb = C_SLATE

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# SECTION 10 — CONSTRAINTS & GOTCHAS
# ════════════════════════════════════════════════════════════════════════════
add_section_banner(doc, '10.  IMPORTANT CONSTRAINTS & KNOWN GOTCHAS',
                   'Engineering notes every developer must know before making changes')

gotchas = [
    ('Sippy Credential Mapping',
     'apiAdminUsername = ssp-root (XML-RPC admin); apiAdminPassword = !chiaan1. '
     'portalUsername = RTST1 (portal login); portalPassword = abcd@1234. '
     'A common mistake is swapping these pairs. The system detects and warns about mismatches. '
     'Always verify which credential pair is required for XML-RPC vs portal HTML scraping.'),
    ('Mixed-Case Usernames Accepted by Sippy',
     'Sippy accepts mixed-case usernames (confirmed: "Internal-PTCL" provisioned successfully). '
     'Do NOT enforce .toLowerCase() on Sippy usernames. '
     'The wizard strips only invalid characters: .replace(/[^a-zA-Z0-9._-]/g, "")'),
    ('Client Accounts Must Be at Root Level',
     'All client accounts must be created at the Sippy root level (i_customer = 1), '
     'NOT under the RTST1 account. Placing an account under RTST1 causes routing '
     'inheritance issues and "No Route Found" errors for accounts without an explicit RG.'),
    ('XML-RPC vs Portal Scraping Fallbacks',
     'Some Sippy versions return HTTP 401 on CDR XML-RPC calls. '
     'The platform falls back to portal HTML scraping in those cases. '
     'Both paths may be active simultaneously — be aware when debugging CDR data.'),
    ('Sippy Load Reduction Guardrails',
     'The platform uses push-based WebSocket NOC, cache-first /api/sippy/live-calls, mutex guards, '
     'and staggered polling intervals to reduce Sippy XML-RPC load by ~65–70 %. '
     'Do not add naive polling loops without these guards — the Sippy instance will be overwhelmed.'),
    ('Auth Middleware Rule',
     'NEVER use the raw isAuthenticated middleware on new routes. '
     'Always use requireRole(["admin", ...], req, res, next) to protect endpoints. '
     'This is the only authorised pattern in this codebase.'),
    ('Rate Card API Availability',
     'On certain Sippy versions, rates can only be added via the Sippy web UI — '
     'no XML-RPC rate creation API is available. '
     'Always verify Sippy version API support before building bulk rate import for a new deployment.'),
    ('Vol 2 Codec & Jitter Features — API Dependency',
     'Vol 2 features #1 (Jitter Monitor) and #2 (Codec Analytics) depend on Sippy exposing '
     'RTP stats and codec fields in its CDR API. '
     'Verify field availability on the target Sippy version before starting implementation.'),
    ('Simulation Mode Disabled in Production',
     'The platform simulation mode is disabled by default (simulationEnabled = false). '
     'Always connect to a live Sippy instance for real data. '
     'Never enable simulation in a production deployment.'),
]

for title, detail in gotchas:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(1)
    r = p.add_run(f'  ⚠  {title}')
    r.bold = True; r.font.size = Pt(10); r.font.color.rgb = C_AMBER
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0); p2.paragraph_format.space_after = Pt(8)
    p2.paragraph_format.left_indent = Cm(0.8)
    r2 = p2.add_run(detail)
    r2.font.size = Pt(9.5); r2.font.color.rgb = C_SLATE

add_rule(doc, C_CYAN)
p_foot = doc.add_paragraph()
p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_foot = p_foot.add_run(
    'BitsAuto Monitoring Platform  \u00b7  Department-Categorised Status & Roadmap Report  \u00b7  May 2026  \u00b7  Confidential'
)
r_foot.font.size = Pt(8.5); r_foot.font.color.rgb = C_MID


# ── Save ─────────────────────────────────────────────────────────────────────
out = '/home/runner/workspace/client/public/downloads/BitsAuto_Platform_Report_May2026.docx'
doc.save(out)
print(f'SUCCESS: {out}')
